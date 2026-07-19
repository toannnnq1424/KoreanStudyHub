#!/usr/bin/env python3
"""Generate the pre-13E Practice Use Case and diagrams.net artifacts.

The generator deliberately keeps one capability model for the DOCX, Draw.io
source and QA previews. Planned nodes are labelled explicitly so an approved
roadmap item cannot be mistaken for code that already exists.
"""

from __future__ import annotations

import argparse
import math
import sys
import textwrap
import xml.etree.ElementTree as ET
from copy import deepcopy
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from practice_use_case_english import (  # noqa: E402
    ACTOR_NAMES,
    BUSINESS_RULES,
    CAPABILITY_COPY,
    DOCUMENT_GROUPS,
    SYSTEM_MESSAGES,
    USE_CASE_COPY,
)


ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = ROOT / "docs" / "architecture" / "practice"
DOCX_PATH = OUTPUT_DIR / "KSH_PRACTICE_USE_CASE_SPECIFICATIONS.docx"
DRAWIO_PATH = OUTPUT_DIR / "KSH_PRACTICE_ARCHITECTURE.drawio.xml"

PAGE_WIDTH = 1600
PAGE_HEIGHT = 1000
DRAWIO_DATE = "2026-07-17T00:00:00.000Z"

COLORS = {
    "ink": "0B2545",
    "blue": "2E74B5",
    "blue_dark": "1F4D78",
    "blue_light": "E8F1F8",
    "muted": "5F6B7A",
    "line": "B8C4D1",
    "panel": "F7F9FC",
    "white": "FFFFFF",
    "current_fill": "D9EAF7",
    "current_stroke": "2E74B5",
    "planned_13e_fill": "FFF2CC",
    "planned_13e_stroke": "D6B656",
    "planned_13f_fill": "FCE8D5",
    "planned_13f_stroke": "D79B00",
    "deferred_13h_fill": "EDE7F6",
    "deferred_13h_stroke": "9673A6",
    "deferred_15_fill": "F8CECC",
    "deferred_15_stroke": "B85450",
    "external_fill": "F2F4F7",
    "external_stroke": "7B8794",
    "success": "2E7D32",
    "warning": "9A6700",
    "error": "B42318",
}

STATUS_STYLE = {
    "CURRENT": (COLORS["current_fill"], COLORS["current_stroke"]),
    "PLANNED 13E": (COLORS["planned_13e_fill"], COLORS["planned_13e_stroke"]),
    "PLANNED 13F": (COLORS["planned_13f_fill"], COLORS["planned_13f_stroke"]),
    "DEFERRED 13H": (COLORS["deferred_13h_fill"], COLORS["deferred_13h_stroke"]),
    "DEFERRED 15": (COLORS["deferred_15_fill"], COLORS["deferred_15_stroke"]),
    "EXTERNAL": (COLORS["external_fill"], COLORS["external_stroke"]),
}


JIRA_DELIVERY_STREAMS: list[dict[str, Any]] = [
    {
        "order": 1,
        "code": "MGT",
        "name": "Practice Test Management",
        "summary": "[Practice][MGT] Deliver Practice Test Management",
        "modules": "AUT, XLS, PDF",
        "depends_on": "None",
        "blocks": "ATT",
        "description": (
            "Provide one governed authoring path for manual editing, Excel import and PDF-assisted import. "
            "Every path must produce the same canonical draft before validation, immutable publication and version management."
        ),
        "acceptance": [
            "A Lecturer can author, validate, publish and manage only Practice content that the Lecturer is authorized to own.",
            "Manual, Excel and PDF input converge into the same canonical draft and publication rules.",
            "Published versions and their referenced assets remain immutable and auditable.",
        ],
        "subtasks": [
            (
                "1",
                "[Practice][MGT][Design] Confirm requirements and Use Case contracts",
                "Approve actors, permissions, lifecycle states and acceptance criteria for UC-AUT-01..03, UC-XLS-01..03 and UC-PDF-01..03.",
            ),
            (
                "2",
                "[Practice][MGT][Design] Maintain the class diagram",
                "Keep one MGT class diagram aligned with current classes, planned boundaries and module ownership.",
            ),
            (
                "3",
                "[Practice][MGT][Design] Maintain nine sequence diagrams",
                "Model each authoring, validation, import, publication and version-management interaction before implementation changes.",
            ),
            (
                "4",
                "[Practice][MGT][Architecture] Finalize draft, asset and version contracts",
                "Define canonical draft mapping, ownership, validation, transaction boundaries, immutable snapshots and failure-safe import behavior.",
            ),
            (
                "5",
                "[Practice][MGT][Backend] Implement authoring and publication services",
                "Implement or refine manual, Excel and PDF adapters so they reuse canonical validation and immutable publication services.",
            ),
            (
                "6",
                "[Practice][MGT][Frontend] Complete the Lecturer workspace",
                "Deliver understandable manual editing, Excel preview and PDF crop/import flows with validation feedback before publication.",
            ),
            (
                "7",
                "[Practice][MGT][Test] Verify security and content integrity",
                "Cover ownership, invalid input, asset access, import atomicity, duplicate submission, publication and historical-version protection.",
            ),
            (
                "8",
                "[Practice][MGT][UAT] Run Lecturer journeys and close evidence",
                "Complete end-to-end authoring journeys, record evidence and resolve all release-blocking linked Bugs.",
            ),
        ],
        "bugs": [
            (
                "Contract",
                "[Practice][MGT][Bug] Import output bypasses canonical draft validation",
                "Blocks Backend and every downstream stream until the reproduced contract breach is fixed.",
            ),
            (
                "Security",
                "[Practice][MGT][Bug] Published version or owned asset can be mutated or exposed incorrectly",
                "Blocks publication and release; link to MGT with the affected version and authorization evidence.",
            ),
            (
                "UX/Data",
                "[Practice][MGT][Bug] Excel or PDF preview differs from the final imported draft",
                "Blocks UAT when the mismatch can change questions, answers, media or scoring data.",
            ),
        ],
    },
    {
        "order": 2,
        "code": "ATT",
        "name": "Skill-based Attempt Lifecycle",
        "summary": "[Practice][ATT] Deliver Skill-based Attempt Lifecycle",
        "modules": "CAT, PLY",
        "depends_on": "MGT publication and version contracts",
        "blocks": "RSL",
        "description": (
            "Provide the Student journey from catalog and preflight through start, resume, autosave, discard and submit. "
            "Reading, Listening, Writing and Speaking share one version-locked attempt lifecycle while retaining skill-specific controls."
        ),
        "acceptance": [
            "A Student sees and starts only an authorized published version.",
            "Timer, device preflight, autosave, resume, discard and submit follow one explicit state machine.",
            "Every saved answer and media item remains bound to the immutable version locked at attempt start.",
        ],
        "subtasks": [
            (
                "1",
                "[Practice][ATT][Design] Confirm lifecycle and preflight contracts",
                "Approve UC-CAT-01..03 and UC-PLY-01..03, including state transitions, timing authority, access and skill-specific device gates.",
            ),
            (
                "2",
                "[Practice][ATT][Design] Maintain the class diagram",
                "Keep the ATT class diagram aligned with catalog, access, version lock, attempt, answer, media and scoring ownership.",
            ),
            (
                "3",
                "[Practice][ATT][Design] Maintain six sequence diagrams",
                "Model browse/detail/start plus save, resume, discard and submit interactions with alternate and failure-safe outcomes.",
            ),
            (
                "4",
                "[Practice][ATT][Architecture] Finalize the attempt state machine",
                "Define version locking, timer authority, idempotency, concurrency, autosave ordering, media ownership and submission boundaries.",
            ),
            (
                "5",
                "[Practice][ATT][Backend] Implement attempt lifecycle services",
                "Implement or refine authorized start, resume, save, discard and submit operations for all four skills.",
            ),
            (
                "6",
                "[Practice][ATT][Frontend] Complete preflight and skill players",
                "Deliver device checks and stable Reading, Listening, Writing and Speaking player states without timer or navigation regressions.",
            ),
            (
                "7",
                "[Practice][ATT][Test] Verify timing, concurrency and media behavior",
                "Cover zero-time starts, duplicate commands, stale autosave, resume, device checks, private speaking media and submission races.",
            ),
            (
                "8",
                "[Practice][ATT][UAT] Run Student attempt journeys",
                "Complete all skill journeys and resolve release-blocking lifecycle Bugs before result work is accepted.",
            ),
        ],
        "bugs": [
            (
                "Lifecycle",
                "[Practice][ATT][Bug] Attempt starts at 00:00 or violates the configured timing authority",
                "Blocks all player UAT and RSL because submitted evidence would be unreliable.",
            ),
            (
                "Concurrency",
                "[Practice][ATT][Bug] Autosave or resume overwrites a newer answer or version lock",
                "Blocks submit and result generation until the reproduced ordering defect has a regression test.",
            ),
            (
                "Device/UX",
                "[Practice][ATT][Bug] Listening or Speaking device check cannot enter the correct player",
                "Blocks the affected skill journey; the check must prove usable playback or recording without unnecessary waiting.",
            ),
        ],
    },
    {
        "order": 3,
        "code": "RSL",
        "name": "Versioned Results and Evidence",
        "summary": "[Practice][RSL] Deliver Versioned Results and Evidence",
        "modules": "RLE, WRT, SPK, RES",
        "depends_on": "ATT submitted attempt and evidence contracts",
        "blocks": "PRG",
        "description": (
            "Present result overview and evidence detail from the exact version used by the attempt. "
            "Objective explanations, Writing evaluation and Speaking evaluation remain separate durable evidence layers with honest pending and failure states."
        ),
        "acceptance": [
            "Result reads never change an answer, score, explanation or evaluation artifact.",
            "Reading and Listening show official answer evidence; Writing and Speaking show skill-appropriate evaluation without fabricated data.",
            "Every overview, detail, retry and review action enforces Student, reviewer and version authorization.",
        ],
        "subtasks": [
            (
                "1",
                "[Practice][RSL][Design] Confirm result and evidence contracts",
                "Approve UC-RLE-01..03, UC-WRT-01..03, UC-SPK-01..03 and UC-RES-01..03, including pending, failed and historical states.",
            ),
            (
                "2",
                "[Practice][RSL][Design] Maintain the class diagram",
                "Keep the RSL class diagram aligned with result assemblers, evidence presenters, durable jobs, artifacts and authorization.",
            ),
            (
                "3",
                "[Practice][RSL][Design] Maintain twelve sequence diagrams",
                "Model objective explanation, Writing and Speaking evaluation, overview, detail and retry or review interactions.",
            ),
            (
                "4",
                "[Practice][RSL][Architecture] Finalize durable evidence contracts",
                "Define fingerprints, queue or outbox lifecycle, idempotency, artifact provenance, retry policy and immutable version reads.",
            ),
            (
                "5",
                "[Practice][RSL][Backend] Implement evidence preparation and result reads",
                "Implement or refine workers, orchestrators, assemblers and presenters without provider calls from result GET requests.",
            ),
            (
                "6",
                "[Practice][RSL][Frontend] Complete overview and evidence detail",
                "Deliver objective, Writing and Speaking result experiences with clear criteria, status, provenance and authorized media playback.",
            ),
            (
                "7",
                "[Practice][RSL][Test] Verify version, provider and authorization safety",
                "Cover wrong-version access, historical attempts, provider timeout, retry, duplicate jobs, partial evidence and reviewer permissions.",
            ),
            (
                "8",
                "[Practice][RSL][UAT] Validate all four result experiences",
                "Compare overview and detail against approved Korean-language scoring criteria and close all release-blocking linked Bugs.",
            ),
        ],
        "bugs": [
            (
                "Version/Auth",
                "[Practice][RSL][Bug] Result exposes evidence from the wrong attempt, version or Student",
                "Security and data-integrity release blocker; fix before any progress aggregation is trusted.",
            ),
            (
                "AI Evidence",
                "[Practice][RSL][Bug] Provider failure fabricates or overwrites an accepted score or explanation",
                "Blocks the affected skill result and retry flow until provenance and idempotency are proven by tests.",
            ),
            (
                "Presentation",
                "[Practice][RSL][Bug] Speaking invents per-question scores or Writing criteria are mislabeled",
                "Blocks UAT because the UI would misrepresent the approved Korean-language scoring contract.",
            ),
        ],
    },
    {
        "order": 4,
        "code": "PRG",
        "name": "Practice Progress Management",
        "summary": "[Practice][PRG] Deliver Practice Progress Management",
        "modules": "PRG",
        "depends_on": "RSL normalized and authorized result contracts",
        "blocks": "Phase 13E exit and Phase 13F recovery acceptance",
        "description": (
            "Provide bounded, authorized progress aggregation across completed Practice attempts. "
            "Students can filter trends, inspect weak areas and open a valid recovery path without mixing incompatible score scales."
        ),
        "acceptance": [
            "Progress includes only authorized attempts and compatible normalized metrics.",
            "Filters, empty states, drill-down and recovery links remain bounded and predictable.",
            "Aggregation and presentation remain read-only and do not trigger evaluation provider work.",
        ],
        "subtasks": [
            (
                "1",
                "[Practice][PRG][Design] Confirm progress and recovery contracts",
                "Approve UC-PRG-01..03, metric definitions, filter behavior, empty states, retention and authorization boundaries.",
            ),
            (
                "2",
                "[Practice][PRG][Design] Maintain the class diagram",
                "Keep the PRG class diagram aligned with bounded queries, aggregation, presentation and recovery-link ownership.",
            ),
            (
                "3",
                "[Practice][PRG][Design] Maintain three sequence diagrams",
                "Model overview, filter or drill-down and recovery interactions with empty, invalid and unauthorized alternatives.",
            ),
            (
                "4",
                "[Practice][PRG][Architecture] Finalize metric and query contracts",
                "Define compatible scales, aggregation windows, pagination, indexes, retention and deep-link authorization.",
            ),
            (
                "5",
                "[Practice][PRG][Backend] Implement bounded progress queries",
                "Implement or refine authorized aggregate, trend, weakness and recovery presenters without unbounded loading.",
            ),
            (
                "6",
                "[Practice][PRG][Frontend] Complete progress and recovery UI",
                "Deliver useful filters, trends, weak-area summaries, empty states, drill-down and recovery commands.",
            ),
            (
                "7",
                "[Practice][PRG][Test] Verify scale, authorization and performance",
                "Cover mixed skills, incompatible scales, no data, large history, invalid filters, cross-Student access and deep links.",
            ),
            (
                "8",
                "[Practice][PRG][UAT] Validate progress decisions",
                "Confirm that displayed insights are understandable, reproducible and linked to valid evidence or recovery journeys.",
            ),
        ],
        "bugs": [
            (
                "Availability/Auth",
                "[Practice][PRG][Bug] Progress query fails or exposes another Student's data",
                "Blocks the entire PRG Task and release; include filter, account and authorization evidence.",
            ),
            (
                "Metrics",
                "[Practice][PRG][Bug] Incompatible score scales are combined into one misleading trend",
                "Blocks the affected metric and recovery recommendation until normalization is corrected.",
            ),
            (
                "Navigation",
                "[Practice][PRG][Bug] Drill-down or recovery opens the wrong attempt or an unauthorized command",
                "Blocks the affected action and must be fixed with deep-link authorization regression coverage.",
            ),
        ],
    },
]


def uc(
    uc_id: str,
    title: str,
    status: str,
    primary: list[str],
    secondary: list[str],
    description: str,
    preconditions: list[str],
    success: list[str],
    failure: list[str],
    steps: list[str],
    alternatives: list[str],
    rules: list[str],
    messages: list[str],
    source: list[str],
    participants: list[str],
    sequence: list[tuple[str, str, str, str]],
) -> dict[str, Any]:
    return {
        "id": uc_id,
        "title": title,
        "status": status,
        "primary": primary,
        "secondary": secondary,
        "description": description,
        "preconditions": preconditions,
        "success": success,
        "failure": failure,
        "steps": steps,
        "alternatives": alternatives,
        "rules": rules,
        "messages": messages,
        "source": source,
        "participants": participants,
        "sequence": sequence,
    }


CAPABILITIES: list[dict[str, Any]] = [
    {
        "code": "CAT",
        "name": "Danh mục, quyền truy cập và điểm vào bài luyện tập",
        "short": "Catalog & Attempt Entry",
        "purpose": "Cho người học tìm đúng bộ đề, xem cấu trúc, rồi bắt đầu hoặc tiếp tục một attempt đã khóa phiên bản bất biến.",
        "actors": ["Người học"],
        "secondary": ["Dịch vụ phân quyền", "Cơ sở dữ liệu"],
        "classes": [
            ("PracticeController", "Controller", "CURRENT", ["catalog/detail routes", "start/resume actions"]),
            ("PracticeCatalogService", "Service", "CURRENT", ["bounded catalog query", "filters and pagination"]),
            ("PracticeDetailPageService", "Service", "CURRENT", ["set/test detail model", "mode/preflight metadata"]),
            ("PracticeLearnerAccessService", "Service", "CURRENT", ["scope and enrolment checks", "deny unpublished content"]),
            ("PracticeService", "Facade", "CURRENT", ["attempt entry orchestration", "player delivery"]),
            ("PracticeAttemptVersionLock", "Value object", "CURRENT", ["published version identity", "immutable delivery boundary"]),
            ("PracticeAttemptRepository", "Repository", "CURRENT", ["active attempt lookup", "attempt persistence"]),
            ("PracticeAttempt", "Entity", "CURRENT", ["status and timestamps", "locked publishedVersionId"]),
        ],
        "relations": [
            ("PracticeController", "PracticeCatalogService", "queries"),
            ("PracticeController", "PracticeDetailPageService", "renders"),
            ("PracticeDetailPageService", "PracticeLearnerAccessService", "authorizes"),
            ("PracticeController", "PracticeService", "starts/resumes"),
            ("PracticeService", "PracticeAttemptRepository", "persists"),
            ("PracticeService", "PracticeAttemptVersionLock", "creates"),
            ("PracticeAttemptRepository", "PracticeAttempt", "stores"),
        ],
        "use_cases": [
            uc(
                "UC-CAT-01", "Duyệt và lọc danh mục luyện tập", "CURRENT",
                ["Người học"], ["Dịch vụ phân quyền"],
                "Hiển thị danh mục bộ đề trong phạm vi người học được phép truy cập, với truy vấn có giới hạn và bộ lọc ổn định.",
                ["Người học đã đăng nhập.", "Tài khoản đang hoạt động."],
                ["Danh mục hợp lệ được hiển thị theo trang.", "Bộ lọc và thứ tự hiển thị được giữ trong URL."],
                ["Không lộ bộ đề nháp, bị khóa hoặc ngoài phạm vi lớp."],
                [
                    "Người học mở trang /practice.",
                    "Hệ thống xác thực phiên và quyền learner.",
                    "Người học chọn bộ lọc kỹ năng, phạm vi hoặc từ khóa.",
                    "Hệ thống chuẩn hóa tham số và áp giới hạn catalog.",
                    "Hệ thống truy vấn các bộ đề người học được phép xem.",
                    "Giao diện hiển thị kết quả, phân trang và trạng thái rỗng phù hợp.",
                ],
                [
                    "A1 - Tham số lọc không hợp lệ: hệ thống dùng giá trị mặc định an toàn và giữ thông báo không gây gián đoạn.",
                    "A2 - Không có kết quả: hệ thống hiển thị empty state thật, không tạo dữ liệu minh họa giả.",
                ],
                ["BR-CAT-01: Catalog phải có giới hạn truy vấn.", "BR-CAT-02: Chỉ nội dung published và đúng scope mới xuất hiện.", "BR-CAT-03: GET danh mục không tạo attempt."],
                ["MSG-CAT-01: Không tìm thấy bộ đề phù hợp.", "MSG-AUTH-01: Phiên đăng nhập đã hết hạn."],
                ["PracticeController", "PracticeCatalogService", "PracticeLearnerAccessService"],
                ["Learner", "Browser/View", "PracticeController", "PracticeCatalogService", "PracticeLearnerAccessService", "Database"],
                [
                    ("Learner", "Browser/View", "open catalog with filters", "call"),
                    ("Browser/View", "PracticeController", "GET /practice", "call"),
                    ("PracticeController", "PracticeCatalogService", "query bounded catalog", "call"),
                    ("PracticeCatalogService", "PracticeLearnerAccessService", "resolve visible scope", "call"),
                    ("PracticeCatalogService", "Database", "fetch permitted page", "call"),
                    ("Database", "PracticeCatalogService", "catalog rows", "return"),
                    ("PracticeCatalogService", "PracticeController", "page model", "return"),
                    ("PracticeController", "Browser/View", "render catalog or empty state", "return"),
                ],
            ),
            uc(
                "UC-CAT-02", "Xem chi tiết bộ đề và bài thi", "CURRENT",
                ["Người học"], ["Dịch vụ phân quyền"],
                "Cho người học kiểm tra cấu trúc set/test/skill, thời lượng và điều kiện preflight trước khi tạo attempt.",
                ["Bộ đề tồn tại.", "Người học có quyền đọc phiên bản đang xuất bản."],
                ["Chi tiết test và các lựa chọn chế độ hợp lệ được hiển thị.", "Không tạo hoặc thay đổi attempt."],
                ["Nội dung lịch sử không thuộc attempt của người học không bị lộ."],
                [
                    "Người học chọn một bộ đề từ catalog.",
                    "Hệ thống kiểm tra quyền theo scope và published version hiện tại.",
                    "Hệ thống nạp danh sách test và kỹ năng theo giới hạn.",
                    "Người học chọn một test để xem chi tiết.",
                    "Hệ thống tính mode, thời lượng và yêu cầu thiết bị.",
                    "Giao diện hiển thị CTA bắt đầu hoặc tiếp tục tương ứng.",
                ],
                [
                    "A1 - Set/test không tồn tại hoặc không được phép: trả trang 404/403 an toàn.",
                    "A2 - Kỹ năng bị khóa: CTA bị vô hiệu và lý do được trình bày rõ.",
                ],
                ["BR-CAT-04: Detail phải dùng published graph hiện tại.", "BR-CAT-05: Historical graph chỉ đọc qua attempt đã khóa phù hợp.", "BR-CAT-06: Preflight được suy ra từ skill delivery."],
                ["MSG-CAT-02: Bạn không có quyền truy cập bài này.", "MSG-CAT-03: Bài luyện tập hiện không khả dụng."],
                ["PracticeController", "PracticeDetailPageService", "PracticeLearnerAccessService"],
                ["Learner", "Browser/View", "PracticeController", "PracticeDetailPageService", "PracticeLearnerAccessService", "Database"],
                [
                    ("Learner", "Browser/View", "select set/test", "call"),
                    ("Browser/View", "PracticeController", "GET set/test detail", "call"),
                    ("PracticeController", "PracticeDetailPageService", "build detail model", "call"),
                    ("PracticeDetailPageService", "PracticeLearnerAccessService", "authorize published graph", "call"),
                    ("PracticeDetailPageService", "Database", "load tests, skills, active attempt", "call"),
                    ("Database", "PracticeDetailPageService", "detail rows", "return"),
                    ("PracticeDetailPageService", "PracticeController", "detail + CTA state", "return"),
                    ("PracticeController", "Browser/View", "render details", "return"),
                ],
            ),
            uc(
                "UC-CAT-03", "Bắt đầu, tiếp tục hoặc hủy attempt", "CURRENT",
                ["Người học"], ["Dịch vụ phân quyền", "Thiết bị trình duyệt"],
                "Điều phối việc tạo attempt mới hoặc tiếp tục attempt đang mở, đồng thời khóa immutable published version và định tuyến preflight đúng kỹ năng.",
                ["Người học được phép làm test.", "Test có published version hợp lệ."],
                ["Attempt có version lock được tạo hoặc tái sử dụng.", "Người học được chuyển tới preflight/player đúng route."],
                ["Không tạo hai active attempt trái hợp đồng.", "Attempt cũ chỉ bị hủy khi người học xác nhận."],
                [
                    "Người học chọn Bắt đầu hoặc Tiếp tục.",
                    "Hệ thống kiểm tra quyền và active attempt hiện có.",
                    "Nếu tạo mới, hệ thống khóa publishedVersionId hiện tại trong transaction.",
                    "Hệ thống tạo snapshot/delivery theo version lock.",
                    "Hệ thống xác định route preflight cho Listening/Speaking hoặc player trực tiếp.",
                    "Trình duyệt chuyển tới route đã xác định.",
                ],
                [
                    "A1 - Có active attempt khác: hiển thị lựa chọn tiếp tục hoặc hủy, không âm thầm ghi đè.",
                    "A2 - Version thay đổi giữa hai bước: transaction thất bại an toàn và yêu cầu tải lại detail.",
                ],
                ["BR-CAT-07: Một attempt luôn khóa đúng một immutable published version.", "BR-CAT-08: Hủy attempt phải có xác nhận và transaction riêng.", "BR-CAT-09: Preflight không được tiêu hao thời gian làm bài."],
                ["MSG-CAT-04: Bạn đang có bài làm chưa hoàn thành.", "MSG-CAT-05: Phiên bản đề đã thay đổi, vui lòng tải lại."],
                ["PracticeService", "PracticeAttemptVersionLock", "PracticeAttemptDiscardService"],
                ["Learner", "Browser/View", "PracticeController", "PracticeService", "PracticeAttemptRepository", "Database"],
                [
                    ("Learner", "Browser/View", "start/resume/discard choice", "call"),
                    ("Browser/View", "PracticeController", "POST attempt action", "call"),
                    ("PracticeController", "PracticeService", "authorize and resolve active attempt", "call"),
                    ("PracticeService", "PracticeAttemptRepository", "lock/find active attempt", "call"),
                    ("PracticeAttemptRepository", "Database", "persist immutable version lock", "call"),
                    ("Database", "PracticeAttemptRepository", "attempt identity", "return"),
                    ("PracticeService", "PracticeController", "preflight/player route", "return"),
                    ("PracticeController", "Browser/View", "redirect", "return"),
                ],
            ),
        ],
    },
    {
        "code": "AUT",
        "name": "Soạn đề thủ công, xuất bản và quản lý phiên bản",
        "short": "Manual Authoring & Publish",
        "purpose": "Cho giảng viên xây dựng graph đề nháp, kiểm tra hợp đồng, xuất bản snapshot bất biến và cộng tác có audit.",
        "actors": ["Giảng viên", "Cộng tác viên"],
        "secondary": ["Quản trị viên", "Kho tài liệu"],
        "classes": [
            ("PracticeDraftController", "Controller", "CURRENT", ["draft editor endpoints", "autosave requests"]),
            ("PracticeManageController", "Controller", "CURRENT", ["manage/publish/revision routes", "feedback states"]),
            ("PracticeDraftService", "Service", "CURRENT", ["mutate draft graph", "edit logging"]),
            ("PracticeDraftContractService", "Service", "CURRENT", ["canonical contract", "skill/type constraints"]),
            ("PracticeDraftValidator", "Validator", "CURRENT", ["publish validation", "actionable violations"]),
            ("PracticePublisherService", "Service", "CURRENT", ["immutable snapshot publish", "after-commit event"]),
            ("PracticeRevisionService", "Service", "CURRENT", ["revision/restore", "published graph safety"]),
            ("PracticeCollaborationService", "Service", "CURRENT", ["collaborator permissions", "audit trail"]),
        ],
        "relations": [
            ("PracticeDraftController", "PracticeDraftService", "edits"),
            ("PracticeDraftService", "PracticeDraftContractService", "applies contract"),
            ("PracticeManageController", "PracticeDraftValidator", "validates"),
            ("PracticeManageController", "PracticePublisherService", "publishes"),
            ("PracticePublisherService", "PracticeRevisionService", "creates baseline"),
            ("PracticeDraftService", "PracticeCollaborationService", "authorizes collaborator"),
        ],
        "use_cases": [
            uc(
                "UC-AUT-01", "Tạo, chỉnh sửa và tự động lưu đề nháp", "CURRENT",
                ["Giảng viên", "Cộng tác viên"], ["Kho tài liệu"],
                "Quản lý graph set/test/section/group/question trong draft với ràng buộc kỹ năng và audit theo người sửa.",
                ["Người dùng có quyền OWNER hoặc COLLABORATOR phù hợp.", "Draft chưa bị khóa bởi trạng thái không cho phép sửa."],
                ["Thay đổi hợp lệ được lưu và ghi audit.", "Preview phản ánh canonical draft mới nhất."],
                ["Mutation trái hợp đồng bị từ chối mà không làm hỏng graph."],
                [
                    "Giảng viên mở editor của một draft.",
                    "Hệ thống kiểm tra quyền chỉnh sửa và trạng thái lifecycle.",
                    "Người dùng thêm hoặc sửa test, section, group hay question.",
                    "Hệ thống chuẩn hóa payload theo assessment contract.",
                    "Hệ thống lưu mutation trong transaction và ghi edit log.",
                    "Giao diện nhận phiên bản draft mới để tiếp tục autosave.",
                ],
                ["A1 - Payload không hợp lệ: trả lỗi tại đúng trường, không lưu một phần.", "A2 - Draft đã xuất bản/khóa: mutation guard chặn và yêu cầu tạo revision."],
                ["BR-AUT-01: Draft mutation phải qua contract service.", "BR-AUT-02: Published graph không được sửa trực tiếp.", "BR-AUT-03: Collaborator action phải được audit."],
                ["MSG-AUT-01: Không thể lưu thay đổi này.", "MSG-AUT-02: Bản đề đã bị khóa để xuất bản."],
                ["PracticeDraftController", "PracticeDraftService", "PracticeDraftContractService"],
                ["Lecturer", "Draft Editor", "PracticeDraftController", "PracticeDraftService", "PracticeDraftContractService", "Database"],
                [
                    ("Lecturer", "Draft Editor", "edit and autosave", "call"),
                    ("Draft Editor", "PracticeDraftController", "submit canonical mutation", "call"),
                    ("PracticeDraftController", "PracticeDraftService", "authorize + mutate", "call"),
                    ("PracticeDraftService", "PracticeDraftContractService", "normalize and enforce rules", "call"),
                    ("PracticeDraftService", "Database", "persist graph + edit log", "call"),
                    ("Database", "PracticeDraftService", "draft version", "return"),
                    ("PracticeDraftService", "PracticeDraftController", "saved draft model", "return"),
                    ("PracticeDraftController", "Draft Editor", "show saved state/errors", "return"),
                ],
            ),
            uc(
                "UC-AUT-02", "Kiểm tra và xuất bản graph đề bất biến", "CURRENT",
                ["Giảng viên"], ["Hàng đợi explanation", "Kho tài liệu"],
                "Kiểm tra toàn bộ draft, tạo immutable version graph trong một transaction và chỉ phát event explanation sau commit.",
                ["Giảng viên là owner có quyền publish.", "Draft tồn tại và storage readiness hợp lệ."],
                ["Published version hoàn chỉnh được tạo.", "Preparation event chỉ xuất hiện sau commit thành công."],
                ["Không công bố graph thiếu hoặc nửa chừng.", "Validation errors vẫn thuộc draft để sửa."],
                [
                    "Giảng viên chọn Kiểm tra trước khi xuất bản.",
                    "Hệ thống chạy validator trên toàn graph và tài liệu tham chiếu.",
                    "Nếu hợp lệ, giảng viên xác nhận xuất bản.",
                    "Publisher khóa draft và sao chép snapshot bất biến trong transaction.",
                    "Transaction commit published version và cập nhật set hiện hành.",
                    "Sau commit, hệ thống phát event chuẩn bị explanation Reading/Listening.",
                ],
                ["A1 - Có violation: hiển thị danh sách theo test/section/question và không publish.", "A2 - Commit thất bại: rollback toàn bộ và không phát event."],
                ["BR-AUT-04: Publish là atomic.", "BR-AUT-05: Snapshot đã publish là immutable.", "BR-AUT-06: AI preparation chạy after-commit, không nằm trong transaction publish."],
                ["MSG-AUT-03: Đề chưa đủ điều kiện xuất bản.", "MSG-AUT-04: Xuất bản thất bại, không có phiên bản nào được công bố."],
                ["PracticeDraftValidator", "PracticePublisherService", "PublishedVersionExplanationEvent"],
                ["Lecturer", "Manage View", "PracticeManageController", "PracticeDraftValidator", "PracticePublisherService", "Database", "After-commit Listener"],
                [
                    ("Lecturer", "Manage View", "validate and confirm publish", "call"),
                    ("Manage View", "PracticeManageController", "POST publish", "call"),
                    ("PracticeManageController", "PracticeDraftValidator", "validate complete graph", "call"),
                    ("PracticeManageController", "PracticePublisherService", "publish immutable snapshot", "call"),
                    ("PracticePublisherService", "Database", "commit graph version", "call"),
                    ("Database", "PracticePublisherService", "published version", "return"),
                    ("PracticePublisherService", "After-commit Listener", "publish explanation event", "call"),
                    ("PracticeManageController", "Manage View", "redirect to published status", "return"),
                ],
            ),
            uc(
                "UC-AUT-03", "Tạo revision, khôi phục và cộng tác", "CURRENT",
                ["Giảng viên", "Cộng tác viên"], ["Quản trị viên"],
                "Tạo draft revision từ version đã xuất bản, phân quyền cộng tác và khôi phục có audit mà không sửa lịch sử bất biến.",
                ["Published version nguồn tồn tại.", "Actor có quyền revision/restore tương ứng."],
                ["Draft revision mới liên kết đúng nguồn.", "Mọi thay đổi quyền và restore có audit."],
                ["Published version cũ không thay đổi.", "Actor ngoài phạm vi không nhận quyền ngầm."],
                [
                    "Owner chọn tạo revision hoặc khôi phục một mốc hợp lệ.",
                    "Hệ thống xác thực quyền và lifecycle của set.",
                    "Revision service sao chép graph nguồn thành draft mới.",
                    "Owner mời hoặc cập nhật quyền collaborator nếu cần.",
                    "Hệ thống ghi nguồn revision, actor và lý do vào audit.",
                    "Editor mở draft mới, còn published version cũ tiếp tục phục vụ attempt lịch sử.",
                ],
                ["A1 - Revision nguồn không còn khả dụng: chặn và giữ nguyên trạng thái.", "A2 - Collaborator thiếu quyền action: trả 403 và ghi audit từ chối khi chính sách yêu cầu."],
                ["BR-AUT-07: Restore tạo draft mới, không mutate published graph.", "BR-AUT-08: Quyền collaborator là explicit.", "BR-AUT-09: Attempt lịch sử tiếp tục dùng version lock cũ."],
                ["MSG-AUT-05: Không thể tạo bản chỉnh sửa từ phiên bản này.", "MSG-AUT-06: Bạn không có quyền thực hiện thao tác."],
                ["PracticeRevisionService", "PracticeCollaborationService", "PracticePublishedGraphMutationGuard"],
                ["Owner/Collaborator", "Manage View", "PracticeManageController", "PracticeRevisionService", "PracticeCollaborationService", "Database"],
                [
                    ("Owner/Collaborator", "Manage View", "request revision/restore", "call"),
                    ("Manage View", "PracticeManageController", "submit action + reason", "call"),
                    ("PracticeManageController", "PracticeRevisionService", "authorize and clone source graph", "call"),
                    ("PracticeRevisionService", "Database", "persist revision + source link", "call"),
                    ("PracticeManageController", "PracticeCollaborationService", "apply collaborator changes", "call"),
                    ("PracticeCollaborationService", "Database", "persist permissions + audit", "call"),
                    ("PracticeManageController", "Manage View", "open new draft", "return"),
                ],
            ),
        ],
    },
    {
        "code": "XLS",
        "name": "Nhập đề bằng Excel",
        "short": "Excel Import",
        "purpose": "Cung cấp template theo contract, preview có lỗi rõ ràng và nhập workbook thành canonical draft có thể chỉnh sửa.",
        "actors": ["Giảng viên"],
        "secondary": ["Kho tài liệu", "Bộ giải mã Excel"],
        "classes": [
            ("PracticeAssessmentExcelController", "Controller", "CURRENT", ["template/preview/import routes", "draft context checks"]),
            ("PracticeAssessmentExcelService", "Service", "CURRENT", ["build template", "preview/import orchestration"]),
            ("PracticeAssessmentExcelV2Codec", "Codec", "CURRENT", ["workbook read/write", "schema validation"]),
            ("AssessmentAuthoringCatalogService", "Service", "CURRENT", ["allowed skills/types", "contract choices"]),
            ("PracticeImportDraftService", "Service", "CURRENT", ["canonical draft materialization", "replace/import boundary"]),
            ("PracticeImportSnapshotService", "Service", "CURRENT", ["source snapshot", "audit/replay evidence"]),
            ("LecturerAssetService", "Service", "CURRENT", ["resolve media links", "ownership checks"]),
            ("PracticeDraftRepository", "Repository", "CURRENT", ["linked draft lookup", "draft persistence"]),
        ],
        "relations": [
            ("PracticeAssessmentExcelController", "PracticeAssessmentExcelService", "delegates"),
            ("PracticeAssessmentExcelService", "PracticeAssessmentExcelV2Codec", "encodes/decodes"),
            ("PracticeAssessmentExcelService", "AssessmentAuthoringCatalogService", "loads contract"),
            ("PracticeAssessmentExcelService", "LecturerAssetService", "resolves media"),
            ("PracticeAssessmentExcelService", "PracticeImportDraftService", "imports"),
            ("PracticeImportDraftService", "PracticeImportSnapshotService", "records source"),
        ],
        "use_cases": [
            uc(
                "UC-XLS-01", "Tải template Excel theo contract", "CURRENT",
                ["Giảng viên"], ["Bộ giải mã Excel"],
                "Sinh workbook có sheet, cột, lựa chọn và dữ liệu mẫu phù hợp với draft context thay vì một template kỹ thuật cố định.",
                ["Giảng viên có draft liên kết hợp lệ.", "Assessment contract của draft đọc được."],
                ["Workbook tải xuống có version/schema rõ ràng.", "Các lựa chọn type/skill phản ánh contract hiện hành."],
                ["Không đưa lựa chọn bị khóa hoặc không hỗ trợ vào template."],
                [
                    "Giảng viên mở workspace Excel từ draft.",
                    "Hệ thống kiểm tra draft và quyền quản lý.",
                    "Hệ thống nạp catalog kỹ năng, dạng câu và giới hạn áp dụng.",
                    "Excel service dựng workbook V2 với hướng dẫn và data validation.",
                    "Codec ghi metadata version cùng các dòng mẫu có ngữ nghĩa.",
                    "Trình duyệt tải file với tên và content type an toàn.",
                ],
                ["A1 - Draft context không hợp lệ: không sinh template và chuyển về manage page.", "A2 - Contract không có kỹ năng khả dụng: hiển thị cấu hình cần sửa."],
                ["BR-XLS-01: Template phải versioned.", "BR-XLS-02: Lựa chọn trong workbook đến từ assessment contract.", "BR-XLS-03: Dòng mẫu phải là dữ liệu học thuật dễ hiểu, không phải placeholder kỹ thuật."],
                ["MSG-XLS-01: Không thể tạo mẫu cho cấu hình hiện tại.", "MSG-XLS-02: Bản nháp liên kết không còn khả dụng."],
                ["PracticeAssessmentExcelController", "PracticeAssessmentExcelService", "PracticeAssessmentExcelV2Codec"],
                ["Lecturer", "Excel Workspace", "PracticeAssessmentExcelController", "PracticeAssessmentExcelService", "AssessmentAuthoringCatalogService", "Excel Codec"],
                [
                    ("Lecturer", "Excel Workspace", "download template", "call"),
                    ("Excel Workspace", "PracticeAssessmentExcelController", "GET template for draft", "call"),
                    ("PracticeAssessmentExcelController", "PracticeAssessmentExcelService", "buildTemplate(context)", "call"),
                    ("PracticeAssessmentExcelService", "AssessmentAuthoringCatalogService", "load allowed contract", "call"),
                    ("PracticeAssessmentExcelService", "Excel Codec", "write versioned workbook", "call"),
                    ("Excel Codec", "PracticeAssessmentExcelService", "workbook bytes", "return"),
                    ("PracticeAssessmentExcelService", "Excel Workspace", "download .xlsx", "return"),
                ],
            ),
            uc(
                "UC-XLS-02", "Xem trước workbook và đối chiếu media", "CURRENT",
                ["Giảng viên"], ["Kho tài liệu", "Bộ giải mã Excel"],
                "Phân tích workbook, hiển thị cấu trúc theo test/skill và lỗi theo dòng trước khi có bất kỳ mutation draft nào.",
                ["File Excel nằm trong giới hạn định dạng/kích thước.", "Draft context còn hiệu lực."],
                ["Preview canonical và danh sách lỗi/cảnh báo được hiển thị.", "Draft chưa bị thay đổi."],
                ["Workbook độc hại, sai version hoặc media ngoài quyền bị từ chối."],
                [
                    "Giảng viên tải workbook lên preview.",
                    "Controller xác thực file, quyền và draft context.",
                    "Codec đọc metadata, sheet và các dòng dữ liệu.",
                    "Service chuẩn hóa thành preview model theo test/skill/group/question.",
                    "Media link được đối chiếu quyền sở hữu và trạng thái.",
                    "Giao diện hiển thị tổng quan, warning và lỗi tại đúng dòng/cột.",
                ],
                ["A1 - Workbook sai schema: trả danh sách lỗi version/sheet/cột.", "A2 - Chỉ có warning: cho phép người dùng quyết định sửa file hoặc tiếp tục import."],
                ["BR-XLS-04: Preview là read-only.", "BR-XLS-05: Lỗi phải định vị được dòng/cột.", "BR-XLS-06: Media không được resolve bằng đường dẫn tùy ý ngoài asset policy."],
                ["MSG-XLS-03: File không đúng mẫu hiện hành.", "MSG-XLS-04: Một số tài nguyên không thuộc thư viện của bạn."],
                ["PracticeAssessmentExcelService", "PracticeAssessmentExcelV2Codec", "LecturerAssetService"],
                ["Lecturer", "Excel Workspace", "PracticeAssessmentExcelController", "PracticeAssessmentExcelService", "Excel Codec", "LecturerAssetService"],
                [
                    ("Lecturer", "Excel Workspace", "upload workbook for preview", "call"),
                    ("Excel Workspace", "PracticeAssessmentExcelController", "POST preview", "call"),
                    ("PracticeAssessmentExcelController", "PracticeAssessmentExcelService", "preview(workbook)", "call"),
                    ("PracticeAssessmentExcelService", "Excel Codec", "decode + validate schema", "call"),
                    ("PracticeAssessmentExcelService", "LecturerAssetService", "resolve referenced media", "call"),
                    ("PracticeAssessmentExcelService", "PracticeAssessmentExcelController", "preview + diagnostics", "return"),
                    ("PracticeAssessmentExcelController", "Excel Workspace", "render structured preview", "return"),
                ],
            ),
            uc(
                "UC-XLS-03", "Nhập workbook thành canonical draft", "CURRENT",
                ["Giảng viên"], ["Kho tài liệu"],
                "Chuyển preview hợp lệ thành graph draft theo một transaction, lưu snapshot nguồn và trả editor có thể chỉnh sửa tiếp.",
                ["Preview hợp lệ hoặc warning đã được xác nhận.", "Draft chưa thay đổi so với context preview."],
                ["Canonical graph được materialize đầy đủ.", "Snapshot nguồn và audit import được lưu."],
                ["Lỗi giữa transaction không để lại graph nửa chừng."],
                [
                    "Giảng viên xác nhận Import từ preview.",
                    "Hệ thống kiểm tra lại quyền, draft version và workbook digest.",
                    "Codec tái giải mã payload thay vì tin dữ liệu client.",
                    "Import service tạo test/section/group/question canonical trong transaction.",
                    "Snapshot service lưu nguồn, diagnostics và actor.",
                    "Giao diện chuyển về editor với bản nháp mới và thông báo kết quả.",
                ],
                ["A1 - Draft đã đổi sau preview: hủy import và yêu cầu preview lại.", "A2 - Transaction thất bại: rollback graph và snapshot, giữ workbook để thử lại an toàn."],
                ["BR-XLS-07: Server phải revalidate khi import.", "BR-XLS-08: Import là atomic.", "BR-XLS-09: Kết quả luôn là canonical draft, không giữ model Excel song song."],
                ["MSG-XLS-05: Bản nháp đã thay đổi, vui lòng xem trước lại.", "MSG-XLS-06: Nhập dữ liệu thất bại; chưa có câu hỏi nào được thêm."],
                ["PracticeImportDraftService", "PracticeImportSnapshotService", "PracticeDraftContractService"],
                ["Lecturer", "Excel Workspace", "PracticeAssessmentExcelController", "PracticeAssessmentExcelService", "PracticeImportDraftService", "Database"],
                [
                    ("Lecturer", "Excel Workspace", "confirm import", "call"),
                    ("Excel Workspace", "PracticeAssessmentExcelController", "POST import", "call"),
                    ("PracticeAssessmentExcelController", "PracticeAssessmentExcelService", "revalidate workbook + context", "call"),
                    ("PracticeAssessmentExcelService", "PracticeImportDraftService", "materialize canonical graph", "call"),
                    ("PracticeImportDraftService", "Database", "persist graph + snapshot atomically", "call"),
                    ("Database", "PracticeImportDraftService", "new draft graph", "return"),
                    ("PracticeImportDraftService", "PracticeAssessmentExcelController", "import summary", "return"),
                    ("PracticeAssessmentExcelController", "Excel Workspace", "redirect to editor", "return"),
                ],
            ),
        ],
    },
    {
        "code": "PDF",
        "name": "Không gian nhập đề từ PDF và AI hỗ trợ",
        "short": "PDF Import Workspace",
        "purpose": "Cho giảng viên xem trang, chọn vùng/crop, kiểm soát evidence gửi AI và biến kết quả thành draft có thể sửa.",
        "actors": ["Giảng viên"],
        "secondary": ["Kho PDF", "Nhà cung cấp AI"],
        "classes": [
            ("PracticeImportController", "Controller", "CURRENT", ["workspace routes", "session entry"]),
            ("PracticePdfImportApiController", "API Controller", "CURRENT", ["page/region/crop/payload APIs", "AI actions"]),
            ("PracticePdfImportSessionService", "Service", "CURRENT", ["session lifecycle", "ownership/status"]),
            ("PracticePdfPageExtractionService", "Service", "CURRENT", ["page images/text", "bounded extraction"]),
            ("PracticePdfRegionService", "Service", "CURRENT", ["region annotations", "stale crop invalidation"]),
            ("PracticePdfCropService", "Service", "CURRENT", ["pixel crop", "asset selection"]),
            ("PracticePdfPayloadPreviewService", "Service", "CURRENT", ["human-readable AI payload", "evidence summary"]),
            ("PracticePdfAiOrchestrator", "Service", "CURRENT", ["provider call/retry", "draft assembly/fallback"]),
        ],
        "relations": [
            ("PracticeImportController", "PracticePdfImportSessionService", "opens"),
            ("PracticePdfImportApiController", "PracticePdfPageExtractionService", "extracts"),
            ("PracticePdfImportApiController", "PracticePdfRegionService", "annotates"),
            ("PracticePdfRegionService", "PracticePdfCropService", "creates crop"),
            ("PracticePdfImportApiController", "PracticePdfPayloadPreviewService", "previews payload"),
            ("PracticePdfImportApiController", "PracticePdfAiOrchestrator", "generates"),
        ],
        "use_cases": [
            uc(
                "UC-PDF-01", "Tạo phiên nhập và kiểm tra trang PDF", "CURRENT",
                ["Giảng viên"], ["Kho PDF"],
                "Tạo import session thuộc owner, trích xuất trang có giới hạn và hiển thị trạng thái xử lý rõ ràng.",
                ["Giảng viên có quyền với draft đích.", "PDF đã qua kiểm tra nội dung và kích thước."],
                ["Session và các page extraction hợp lệ được tạo.", "Workspace hiển thị thumbnail/text có trạng thái."],
                ["File hỏng hoặc ngoài quyền không tạo session sử dụng được."],
                [
                    "Giảng viên chọn tài liệu PDF trong thư viện.",
                    "Hệ thống xác thực quyền sở hữu và draft context.",
                    "Session service tạo phiên import có owner và trạng thái ban đầu.",
                    "Page extraction service xử lý số trang trong giới hạn.",
                    "Hệ thống lưu ảnh/text extraction cùng checksum.",
                    "Workspace hiển thị danh sách trang và lỗi từng trang nếu có.",
                ],
                ["A1 - Một trang trích xuất lỗi: đánh dấu trang đó và cho phép phần còn lại tiếp tục khi an toàn.", "A2 - File đã thay đổi: vô hiệu session cũ và yêu cầu tạo lại."],
                ["BR-PDF-01: Session thuộc một owner và draft context.", "BR-PDF-02: Page extraction có giới hạn tài nguyên.", "BR-PDF-03: Không dùng đường dẫn file client làm authority."],
                ["MSG-PDF-01: Không thể đọc tài liệu PDF.", "MSG-PDF-02: Phiên nhập đã hết hiệu lực."],
                ["PracticePdfImportSessionService", "PracticePdfPageExtractionService", "PracticePdfStorageService"],
                ["Lecturer", "PDF Workspace", "PracticeImportController", "PracticePdfImportSessionService", "PracticePdfPageExtractionService", "Database"],
                [
                    ("Lecturer", "PDF Workspace", "select PDF", "call"),
                    ("PDF Workspace", "PracticeImportController", "create import session", "call"),
                    ("PracticeImportController", "PracticePdfImportSessionService", "authorize + create", "call"),
                    ("PracticePdfImportSessionService", "PracticePdfPageExtractionService", "extract bounded pages", "call"),
                    ("PracticePdfPageExtractionService", "Database", "persist page evidence", "call"),
                    ("Database", "PracticePdfImportSessionService", "session + pages", "return"),
                    ("PracticeImportController", "PDF Workspace", "render workspace", "return"),
                ],
            ),
            uc(
                "UC-PDF-02", "Đánh dấu vùng, crop ảnh và xem payload AI", "CURRENT",
                ["Giảng viên"], ["Kho PDF"],
                "Cho giảng viên chọn vùng nội dung, crop ảnh và xem chính xác evidence nào sẽ được gửi AI trước khi phát sinh chi phí.",
                ["Import session đang hoạt động.", "Page extraction mục tiêu tồn tại."],
                ["Region/crop hợp lệ được lưu.", "Payload preview phản ánh đúng vùng đang chọn."],
                ["Crop cũ bị vô hiệu khi region thay đổi.", "Không gửi AI trong bước preview."],
                [
                    "Giảng viên chọn trang và vẽ vùng nội dung.",
                    "API chuẩn hóa tọa độ theo kích thước trang thực.",
                    "Region service lưu annotation và vô hiệu crop stale.",
                    "Crop service tạo ảnh crop trong ranh giới an toàn.",
                    "Giảng viên gán loại vùng và nhóm câu hỏi.",
                    "Payload preview hiển thị text, ảnh, chiến lược và phần bị loại trừ.",
                ],
                ["A1 - Tọa độ ngoài trang: từ chối và giữ region cũ.", "A2 - Crop không tạo được: cho sửa region hoặc tiếp tục bằng text khi policy cho phép."],
                ["BR-PDF-04: UI phải có thao tác crop nhìn thấy được.", "BR-PDF-05: Region đổi thì crop cũ không còn hợp lệ.", "BR-PDF-06: Preview payload là read-only và không gọi provider."],
                ["MSG-PDF-03: Vùng chọn không hợp lệ.", "MSG-PDF-04: Ảnh cắt đã cũ, vui lòng tạo lại."],
                ["PracticePdfRegionService", "PracticePdfCropService", "PracticePdfPayloadPreviewService"],
                ["Lecturer", "PDF Workspace", "PracticePdfImportApiController", "PracticePdfRegionService", "PracticePdfCropService", "PracticePdfPayloadPreviewService"],
                [
                    ("Lecturer", "PDF Workspace", "draw region and request crop", "call"),
                    ("PDF Workspace", "PracticePdfImportApiController", "save annotation", "call"),
                    ("PracticePdfImportApiController", "PracticePdfRegionService", "normalize + persist region", "call"),
                    ("PracticePdfRegionService", "PracticePdfCropService", "invalidate old and create crop", "call"),
                    ("PDF Workspace", "PracticePdfImportApiController", "preview AI payload", "call"),
                    ("PracticePdfImportApiController", "PracticePdfPayloadPreviewService", "build readable preview", "call"),
                    ("PracticePdfPayloadPreviewService", "PDF Workspace", "text/image evidence summary", "return"),
                ],
            ),
            uc(
                "UC-PDF-03", "Sinh draft bằng AI, thử lại hoặc chuyển nhập tay", "CURRENT",
                ["Giảng viên"], ["Nhà cung cấp AI"],
                "Gọi AI chỉ sau xác nhận payload, kiểm tra kết quả và lắp thành canonical draft; luôn có retry có kiểm soát và manual fallback.",
                ["Payload preview đã được xác nhận.", "Session/evidence chưa stale.", "AI configuration sẵn sàng hoặc manual fallback được chọn."],
                ["Kết quả hợp lệ được lắp thành draft để giảng viên sửa.", "Provider audit và snapshot được lưu."],
                ["Provider lỗi không tạo điểm/câu hỏi giả.", "Payload không hợp lệ không mutate draft."],
                [
                    "Giảng viên xác nhận gửi payload đã xem trước.",
                    "Orchestrator xây payload server-side và ghi audit request.",
                    "AI client gửi request có timeout và giới hạn.",
                    "Kết quả được validate rồi assembler chuyển thành import model.",
                    "Import service materialize canonical draft trong transaction.",
                    "Giao diện hiển thị draft, diagnostics và lựa chọn chỉnh tay.",
                ],
                ["A1 - Provider timeout/không hợp lệ: giữ session, hiển thị retry có lý do.", "A2 - AI không sẵn sàng: cho nhập tay từ vùng đã chọn mà không mất evidence."],
                ["BR-PDF-07: Chỉ server xây provider payload cuối cùng.", "BR-PDF-08: Provider failure không fabricate content.", "BR-PDF-09: AI output phải qua validator và canonical importer."],
                ["MSG-PDF-05: AI chưa thể xử lý tài liệu này.", "MSG-PDF-06: Dữ liệu sinh ra cần được sửa trước khi nhập."],
                ["PracticePdfAiOrchestrator", "PracticePdfDraftAssembler", "PracticeImportDraftService"],
                ["Lecturer", "PDF Workspace", "PracticePdfImportApiController", "PracticePdfAiOrchestrator", "AI Provider", "PracticeImportDraftService", "Database"],
                [
                    ("Lecturer", "PDF Workspace", "confirm AI generation", "call"),
                    ("PDF Workspace", "PracticePdfImportApiController", "POST generate", "call"),
                    ("PracticePdfImportApiController", "PracticePdfAiOrchestrator", "build audited request", "call"),
                    ("PracticePdfAiOrchestrator", "AI Provider", "generate structured questions", "call"),
                    ("AI Provider", "PracticePdfAiOrchestrator", "result or typed failure", "return"),
                    ("PracticePdfAiOrchestrator", "PracticeImportDraftService", "validate + materialize", "call"),
                    ("PracticeImportDraftService", "Database", "persist canonical draft", "call"),
                    ("PracticePdfImportApiController", "PDF Workspace", "show draft/retry/manual fallback", "return"),
                ],
            ),
        ],
    },
    {
        "code": "PLY",
        "name": "Player theo kỹ năng và vòng đời attempt",
        "short": "Skill-native Player",
        "purpose": "Cung cấp preflight, player, autosave/resume, submit/timeout/discard dựa trên immutable attempt delivery.",
        "actors": ["Người học"],
        "secondary": ["Thiết bị âm thanh/micro", "Kho media riêng tư"],
        "classes": [
            ("PracticeController", "Controller", "CURRENT", ["preflight/player routes", "save/submit/discard"]),
            ("PracticeService", "Facade", "CURRENT", ["immutable section delivery", "autosave/submit orchestration"]),
            ("PracticeVersionSnapshot", "Value object", "CURRENT", ["attempt graph snapshot", "delivery consistency"]),
            ("PracticeAnswerFormMapper", "Mapper", "CURRENT", ["request to learner answers", "canonical field mapping"]),
            ("AssessmentScoringEngine", "Domain service", "CURRENT", ["objective scoring", "partial/essay status"]),
            ("PracticeAttemptDiscardService", "Service", "CURRENT", ["confirmed discard", "transaction boundary"]),
            ("PracticeSpeakingMediaService", "Service", "CURRENT", ["private recording lifecycle", "attempt ownership"]),
            ("PracticeAttemptRepository", "Repository", "CURRENT", ["attempt state/answers", "row locking"]),
        ],
        "relations": [
            ("PracticeController", "PracticeService", "delivers/saves/submits"),
            ("PracticeService", "PracticeVersionSnapshot", "loads"),
            ("PracticeService", "PracticeAnswerFormMapper", "maps"),
            ("PracticeService", "AssessmentScoringEngine", "scores"),
            ("PracticeController", "PracticeAttemptDiscardService", "discards"),
            ("PracticeService", "PracticeSpeakingMediaService", "activates media"),
            ("PracticeService", "PracticeAttemptRepository", "locks/persists"),
        ],
        "use_cases": [
            uc(
                "UC-PLY-01", "Kiểm tra thiết bị trước bài Nghe/Nói", "CURRENT",
                ["Người học"], ["Thiết bị âm thanh/micro"],
                "Xác nhận trình duyệt có thể bắt đầu phát mẫu nghe hoặc cấp quyền thu âm trước khi vào player, không tiêu hao thời gian thi.",
                ["Attempt đã được tạo và đang ở trạng thái cho phép preflight.", "Immutable delivery có media/preflight metadata hợp lệ."],
                ["Thiết bị được xác nhận và learner có thể vào player.", "Timer làm bài chỉ bắt đầu trong player theo server time."],
                ["Thiếu media hoặc lỗi thiết bị được trình bày rõ; không kẹt im lặng ở cùng trang."],
                [
                    "Người học mở route preflight của attempt.",
                    "Hệ thống nạp preflight từ immutable attempt version.",
                    "Người học nhấn kiểm tra âm thanh hoặc micro.",
                    "Trình duyệt bắt đầu playback/permission request và ghi nhận kết quả.",
                    "Khi playback bắt đầu nghe được hoặc micro sẵn sàng, CTA vào bài được bật.",
                    "Người học tiếp tục và trình duyệt chuyển tới player.",
                ],
                ["A1 - Media không tồn tại/không phát được: hiển thị lỗi và đường thử lại, không coi là pass.", "A2 - Người dùng từ chối micro: hướng dẫn cấp lại quyền, không tự khởi động attempt timer."],
                ["BR-PLY-01: Listening check chỉ cần playback bắt đầu, không đợi audio kết thúc.", "BR-PLY-02: Preflight không làm timer về 00:00.", "BR-PLY-03: Route tiếp tục phải có đích rõ ràng sau khi check thành công."],
                ["MSG-PLY-01: Chưa thể phát âm thanh mẫu.", "MSG-PLY-02: Trình duyệt chưa được cấp quyền micro."],
                ["PracticeController", "PracticeService.getAttemptListeningPreflightDelivery", "PracticeMediaRoutes"],
                ["Learner", "Preflight View", "Browser Media API", "PracticeController", "PracticeService", "Database"],
                [
                    ("Learner", "Preflight View", "open device check", "call"),
                    ("Preflight View", "PracticeController", "GET immutable preflight", "call"),
                    ("PracticeController", "PracticeService", "load attempt preflight delivery", "call"),
                    ("PracticeService", "Database", "read locked media reference", "call"),
                    ("PracticeController", "Preflight View", "render sample + player target", "return"),
                    ("Preflight View", "Browser Media API", "start playback/permission", "call"),
                    ("Browser Media API", "Preflight View", "playing/ready or error", "return"),
                    ("Preflight View", "Learner", "enable continue or show recovery", "return"),
                ],
            ),
            uc(
                "UC-PLY-02", "Làm bài, tự động lưu và tiếp tục", "CURRENT",
                ["Người học"], ["Kho media riêng tư"],
                "Hiển thị player theo kỹ năng từ immutable snapshot, lưu câu trả lời có xung đột rõ ràng và khôi phục timer/position khi resume.",
                ["Attempt đang IN_PROGRESS.", "Server còn thời gian hoặc policy cho phép khôi phục."],
                ["Câu trả lời mới nhất được lưu.", "Resume dùng cùng version, timer và vị trí hợp lệ."],
                ["Client cũ không ghi đè dữ liệu mới hơn.", "Không lộ key hoặc explanation trong player payload."],
                [
                    "Người học mở player của attempt.",
                    "Hệ thống nạp version snapshot và section delivery theo skill.",
                    "Giao diện dựng câu hỏi, media, timer và navigation cố định.",
                    "Người học trả lời; client gửi autosave theo nhịp hoặc thao tác.",
                    "Server khóa attempt, kiểm tra version/state rồi lưu canonical answers.",
                    "Khi tải lại, player phục hồi answer, vị trí và thời gian từ server contract.",
                ],
                ["A1 - Autosave xung đột: trả conflict và yêu cầu đồng bộ thay vì ghi đè.", "A2 - Mạng lỗi: giữ trạng thái chưa lưu nhìn thấy được và thử lại có giới hạn."],
                ["BR-PLY-04: Player payload không chứa official key.", "BR-PLY-05: Server time là authority cho deadline.", "BR-PLY-06: Autosave phải idempotent theo attempt state/version."],
                ["MSG-PLY-03: Có thay đổi chưa được lưu.", "MSG-PLY-04: Bài làm đã thay đổi ở phiên khác."],
                ["PracticeService", "PracticeVersionSnapshot", "PracticeAnswerFormMapper"],
                ["Learner", "Skill Player", "PracticeController", "PracticeService", "PracticeAttemptRepository", "Database"],
                [
                    ("Learner", "Skill Player", "open/resume player", "call"),
                    ("Skill Player", "PracticeController", "GET attempt player", "call"),
                    ("PracticeController", "PracticeService", "load immutable delivery", "call"),
                    ("PracticeService", "PracticeAttemptRepository", "load state + answers", "call"),
                    ("PracticeController", "Skill Player", "render delivery without key", "return"),
                    ("Skill Player", "PracticeController", "POST autosave", "call"),
                    ("PracticeController", "PracticeService", "validate + save canonical answers", "call"),
                    ("PracticeService", "Database", "persist answers/deadline state", "call"),
                    ("PracticeController", "Skill Player", "saved/conflict state", "return"),
                ],
            ),
            uc(
                "UC-PLY-03", "Nộp bài, hết giờ hoặc hủy attempt", "CURRENT",
                ["Người học"], ["Dịch vụ chấm điểm"],
                "Kết thúc attempt đúng một lần qua submit hoặc timeout, chấm theo immutable spec và tách hủy có xác nhận khỏi nộp bài.",
                ["Attempt tồn tại và thuộc người học.", "State transition được phép."],
                ["Attempt chuyển trạng thái kết thúc duy nhất.", "Objective score và AI status được lưu đúng contract."],
                ["Không double-submit, không fabricated score và không biến discard thành submit."],
                [
                    "Người học chọn Nộp bài hoặc deadline được xác nhận đã hết.",
                    "Controller gửi canonical answer payload và action type.",
                    "Service khóa attempt và kiểm tra state/deadline server-side.",
                    "Scoring engine chấm objective answers từ immutable answer spec.",
                    "Writing/Speaking được chuyển sang evaluation status tương ứng.",
                    "Transaction lưu completion và trả route result phù hợp.",
                ],
                ["A1 - Submit lặp: trả kết quả hiện hữu, không chấm lại ngầm.", "A2 - Người học chọn Hủy: chạy discard transaction riêng sau xác nhận và không tạo result giả."],
                ["BR-PLY-07: State transition phải atomic.", "BR-PLY-08: Objective scoring dùng immutable answer spec.", "BR-PLY-09: Provider failure không sinh score mặc định."],
                ["MSG-PLY-05: Bài đã được nộp trước đó.", "MSG-PLY-06: Không thể hủy bài ở trạng thái hiện tại."],
                ["PracticeService.submitAttempt", "AssessmentScoringEngine", "PracticeAttemptDiscardService"],
                ["Learner", "Skill Player", "PracticeController", "PracticeService", "AssessmentScoringEngine", "Database"],
                [
                    ("Learner", "Skill Player", "submit / confirm discard", "call"),
                    ("Skill Player", "PracticeController", "POST terminal action", "call"),
                    ("PracticeController", "PracticeService", "lock + validate state", "call"),
                    ("PracticeService", "AssessmentScoringEngine", "score immutable objective answers", "call"),
                    ("AssessmentScoringEngine", "PracticeService", "score/status", "return"),
                    ("PracticeService", "Database", "persist terminal state atomically", "call"),
                    ("Database", "PracticeService", "completed/discarded attempt", "return"),
                    ("PracticeController", "Skill Player", "redirect result or catalog", "return"),
                ],
            ),
        ],
    },
    {
        "code": "RLE",
        "name": "Vòng đời giải thích bất biến cho Đọc/Nghe",
        "short": "R/L Explanation Lifecycle",
        "purpose": "Chuẩn bị, tái sử dụng và đọc explanation theo fingerprint nội dung sau publish mà không gọi AI trong result GET.",
        "actors": ["Giảng viên", "Người vận hành", "Người học"],
        "secondary": ["Nhà cung cấp AI", "Worker"],
        "classes": [
            ("PublishedVersionExplanationListener", "Listener", "CURRENT", ["after-commit intake", "prepare published version"]),
            ("QuestionExplanationPreparationService", "Service", "CURRENT", ["eligible question scan", "binding/task preparation"]),
            ("ExplanationInputFactory", "Factory", "CURRENT", ["immutable prompt/evidence", "language-safe input"]),
            ("ExplanationFingerprintBuilder", "Service", "CURRENT", ["ID-independent fingerprint", "content identity"]),
            ("QuestionExplanationGenerationWorker", "Worker", "CURRENT", ["claim durable tasks", "bounded processing"]),
            ("QuestionExplanationRetryService", "Service", "CURRENT", ["explicit retry", "status policy"]),
            ("QuestionExplanationReadService", "Service", "CURRENT", ["binding-safe read", "no provider calls"]),
            ("ObjectiveEvidencePresenter", "Logical presenter", "PLANNED 13E", ["official/learner/teacher/AI layers", "evidence anchors"]),
        ],
        "relations": [
            ("PublishedVersionExplanationListener", "QuestionExplanationPreparationService", "prepares"),
            ("QuestionExplanationPreparationService", "ExplanationInputFactory", "builds input"),
            ("QuestionExplanationPreparationService", "ExplanationFingerprintBuilder", "identifies"),
            ("QuestionExplanationGenerationWorker", "QuestionExplanationRetryService", "shares task policy"),
            ("QuestionExplanationGenerationWorker", "QuestionExplanationReadService", "produces readable artifact"),
            ("QuestionExplanationReadService", "ObjectiveEvidencePresenter", "will feed"),
        ],
        "use_cases": [
            uc(
                "UC-RLE-01", "Chuẩn bị artifact sau khi xuất bản", "CURRENT",
                ["Giảng viên"], ["After-commit listener", "Worker"],
                "Sau publish commit, xác định câu Đọc/Nghe đủ điều kiện, tạo input bất biến, fingerprint, binding và durable task.",
                ["Published version đã commit thành công.", "Question version và evidence tham chiếu đọc được."],
                ["Mỗi question version có binding tới artifact identity.", "Task được tạo hoặc reuse idempotently."],
                ["Publish không rollback vì provider.", "Crash gap có thể được reconciler bù lại."],
                [
                    "Publisher phát event sau commit.",
                    "Listener gọi preparation service với publishedVersionId.",
                    "Service quét các Reading/Listening question version đủ điều kiện.",
                    "Input factory dựng prompt, options, key và evidence bất biến.",
                    "Fingerprint builder tính identity không phụ thuộc live ID.",
                    "Transaction tạo/reuse artifact, binding và generation task.",
                ],
                ["A1 - Question không đủ evidence: ghi trạng thái UNAVAILABLE có lý do, không gọi AI.", "A2 - Listener bị gián đoạn: reconciler phát hiện published version chưa chuẩn bị và chạy bù idempotent."],
                ["BR-RLE-01: Event chỉ chạy after commit.", "BR-RLE-02: Fingerprint không phụ thuộc database ID.", "BR-RLE-03: Binding luôn explicit theo immutable question version."],
                ["MSG-RLE-01: Explanation chưa có đủ dữ liệu nguồn.", "MSG-RLE-02: Preparation đang chờ xử lý lại."],
                ["PublishedVersionExplanationListener", "QuestionExplanationPreparationService", "ExplanationFingerprintBuilder"],
                ["Publisher", "After-commit Listener", "PreparationService", "InputFactory", "FingerprintBuilder", "Database"],
                [
                    ("Publisher", "After-commit Listener", "published version event", "call"),
                    ("After-commit Listener", "PreparationService", "prepare(versionId)", "call"),
                    ("PreparationService", "InputFactory", "build immutable inputs", "call"),
                    ("PreparationService", "FingerprintBuilder", "compute content identity", "call"),
                    ("PreparationService", "Database", "upsert artifact/binding/task", "call"),
                    ("Database", "PreparationService", "prepared counts/status", "return"),
                    ("PreparationService", "After-commit Listener", "preparation result", "return"),
                ],
            ),
            uc(
                "UC-RLE-02", "Sinh, tái sử dụng và thử lại explanation", "CURRENT",
                ["Worker", "Người vận hành"], ["Nhà cung cấp AI"],
                "Worker claim durable task, reuse READY artifact theo fingerprint hoặc gọi provider có audit; retry là thao tác có trạng thái rõ ràng.",
                ["Task ở trạng thái có thể claim.", "Provider readiness cho phép gọi hoặc artifact READY đã tồn tại."],
                ["Artifact READY hợp lệ được lưu/reuse.", "Task terminal/retryable được cập nhật nhất quán."],
                ["Không tạo artifact READY từ output lỗi.", "Retry không tạo duplicate identity."],
                [
                    "Worker claim một task bằng transaction ngắn.",
                    "Processor nạp immutable work và kiểm tra artifact hiện tại.",
                    "Nếu fingerprint đã READY, task hoàn tất bằng reuse.",
                    "Nếu chưa có, client gọi provider với timeout/audit.",
                    "Output được validate và lưu artifact READY trong transaction.",
                    "Lỗi được phân loại retryable/terminal và hiển thị cho operator.",
                ],
                ["A1 - Provider timeout: lên lịch retry theo policy và không đổi artifact thành READY.", "A2 - Operator retry: service kiểm tra trạng thái và reset task có audit, không gọi provider đồng bộ trong request."],
                ["BR-RLE-04: Worker claim và provider call không giữ transaction dài.", "BR-RLE-05: READY artifact là immutable.", "BR-RLE-06: Retry là durable state transition."],
                ["MSG-RLE-03: Explanation đang được tạo.", "MSG-RLE-04: Tạo explanation thất bại; có thể thử lại."],
                ["QuestionExplanationGenerationWorker", "QuestionExplanationGenerationProcessor", "QuestionExplanationRetryService"],
                ["Worker/Operator", "Task Worker", "GenerationProcessor", "WorkLoader", "AI Provider", "TaskTransactions", "Database"],
                [
                    ("Worker/Operator", "Task Worker", "claim or request retry", "call"),
                    ("Task Worker", "TaskTransactions", "claim durable task", "call"),
                    ("Task Worker", "GenerationProcessor", "process claimed work", "call"),
                    ("GenerationProcessor", "WorkLoader", "load immutable input", "call"),
                    ("GenerationProcessor", "AI Provider", "generate when no READY reuse", "call"),
                    ("AI Provider", "GenerationProcessor", "validated result or typed error", "return"),
                    ("GenerationProcessor", "TaskTransactions", "persist READY/failure state", "call"),
                    ("TaskTransactions", "Database", "commit artifact/task", "call"),
                ],
            ),
            uc(
                "UC-RLE-03", "Đọc explanation bất biến trong kết quả", "CURRENT",
                ["Người học"], ["Result assembler"],
                "Tra explanation qua binding của immutable question version và trả status an toàn; đường GET tuyệt đối không gọi provider.",
                ["Người học sở hữu attempt hoặc có quyền review.", "Attempt khóa published version hợp lệ."],
                ["Artifact/binding phù hợp được đưa vào result model.", "Learner overlay tách khỏi shared artifact."],
                ["Không lộ artifact của version khác.", "Pending/failed/unavailable không bị giả thành explanation thành công."],
                [
                    "Người học mở result hoặc detail của attempt.",
                    "Hệ thống xác thực attempt ownership/reviewer permission.",
                    "Result context nạp immutable question versions của attempt.",
                    "Read service tra binding theo từng question version.",
                    "Service trả READY artifact hoặc status pending/failed/unavailable.",
                    "Presenter ghép learner answer overlay riêng mà không sửa artifact.",
                ],
                ["A1 - Không có binding: trả unavailable có lý do và không fallback sang version hiện tại.", "A2 - Artifact pending/failed: hiển thị trạng thái/recovery phù hợp, không gọi AI trong GET."],
                ["BR-RLE-07: Result GET là read-only.", "BR-RLE-08: Lookup theo immutable question version binding.", "BR-RLE-09: Learner overlay không thuộc shared artifact."],
                ["MSG-RLE-05: Giải thích đang được chuẩn bị.", "MSG-RLE-06: Giải thích chưa khả dụng cho phiên bản câu hỏi này."],
                ["QuestionExplanationReadService", "PracticeResultAssembler", "ObjectiveResultPresenter"],
                ["Learner", "Result View", "PracticeController", "PracticeResultAssembler", "QuestionExplanationReadService", "Database"],
                [
                    ("Learner", "Result View", "open result/detail", "call"),
                    ("Result View", "PracticeController", "GET attempt result", "call"),
                    ("PracticeController", "PracticeResultAssembler", "assemble immutable result", "call"),
                    ("PracticeResultAssembler", "QuestionExplanationReadService", "read by question-version binding", "call"),
                    ("QuestionExplanationReadService", "Database", "fetch binding + artifact status", "call"),
                    ("Database", "QuestionExplanationReadService", "READY/pending/failed/unavailable", "return"),
                    ("PracticeResultAssembler", "Result View", "render artifact + separate overlay", "return"),
                ],
            ),
        ],
    },
    {
        "code": "WRT",
        "name": "Chấm Viết bằng AI theo tiêu chí tiếng Hàn",
        "short": "Writing AI Evaluation",
        "purpose": "Chấm bài tự luận bằng rubric tiếng Hàn, chuẩn hóa evidence và chỉ tái chấm khi người dùng có quyền yêu cầu.",
        "actors": ["Người học", "Người đánh giá"],
        "secondary": ["Nhà cung cấp AI"],
        "classes": [
            ("WritingEvaluationCacheService", "Service", "CURRENT", ["identity/reuse", "persist evaluation status"]),
            ("WritingEvaluationClient", "Provider port", "CURRENT", ["provider request", "typed failure"]),
            ("WritingEvaluationNormalizer", "Normalizer", "CURRENT", ["schema/score normalization", "reject malformed feedback"]),
            ("WritingPromptRules", "Rules", "CURRENT", ["Korean task prompt", "safe evidence boundaries"]),
            ("WritingRuleEngine", "Domain service", "CURRENT", ["rubric invariants", "criterion checks"]),
            ("WritingScoringPolicy", "Domain service", "CURRENT", ["TOPIK weights", "score aggregation"]),
            ("WritingResultPresenter", "Presenter", "CURRENT", ["skill-native overview", "four analysis lenses"]),
            ("WritingEvidencePresenter", "Logical presenter", "PLANNED 13E", ["submitted/correction/rewrite/sample", "rubric evidence detail"]),
        ],
        "relations": [
            ("WritingEvaluationCacheService", "WritingEvaluationClient", "calls when not reusable"),
            ("WritingEvaluationClient", "WritingEvaluationNormalizer", "returns provider data"),
            ("WritingEvaluationNormalizer", "WritingRuleEngine", "validates"),
            ("WritingRuleEngine", "WritingScoringPolicy", "scores"),
            ("WritingEvaluationCacheService", "WritingResultPresenter", "feeds"),
            ("WritingResultPresenter", "WritingEvidencePresenter", "will share context"),
        ],
        "use_cases": [
            uc(
                "UC-WRT-01", "Chấm bài viết đã nộp", "CURRENT",
                ["Người học"], ["Nhà cung cấp AI"],
                "Tạo yêu cầu chấm từ immutable question/task và learner answer, chuẩn hóa theo rubric tiếng Hàn rồi lưu score/status có provenance.",
                ["Attempt Writing đã submit.", "Câu hỏi là ESSAY hợp lệ và có answer text."],
                ["Evaluation hợp lệ có score, rubric evidence và provenance.", "Result không hiển thị score giả khi provider lỗi."],
                ["Malformed feedback bị từ chối.", "Answer rỗng xử lý theo rule thay vì provider tùy ý."],
                [
                    "Submit service xác định các Writing answers cần chấm.",
                    "Evaluation service dựng identity từ immutable task và answer.",
                    "Prompt rules tạo yêu cầu tiếng Hàn theo loại câu 51-54.",
                    "Client gọi provider có audit/timeout.",
                    "Normalizer kiểm tra schema, rubric và score bounds.",
                    "Rule engine/scoring policy lưu kết quả hoặc typed failure.",
                ],
                ["A1 - Answer thiếu/không đủ điều kiện: lưu trạng thái theo rule, không gọi provider vô nghĩa.", "A2 - Provider lỗi: lưu failure/pending phù hợp và không tạo score mặc định."],
                ["BR-WRT-01: Q51-Q54 production canonical là ESSAY trong scope rút gọn.", "BR-WRT-02: Score phải khớp rubric tiếng Hàn và trọng số task.", "BR-WRT-03: Provider failure không fabricate score."],
                ["MSG-WRT-01: Bài viết đang được chấm.", "MSG-WRT-02: Chưa thể chấm bài viết; vui lòng thử lại sau."],
                ["WritingEvaluationCacheService", "WritingEvaluationClient", "WritingEvaluationNormalizer"],
                ["Submit Service", "WritingEvaluationCacheService", "WritingPromptRules", "WritingEvaluationClient", "AI Provider", "WritingEvaluationNormalizer", "Database"],
                [
                    ("Submit Service", "WritingEvaluationCacheService", "evaluate immutable answer", "call"),
                    ("WritingEvaluationCacheService", "WritingPromptRules", "build Korean rubric request", "call"),
                    ("WritingEvaluationCacheService", "WritingEvaluationClient", "request evaluation", "call"),
                    ("WritingEvaluationClient", "AI Provider", "provider call", "call"),
                    ("AI Provider", "WritingEvaluationClient", "feedback or typed error", "return"),
                    ("WritingEvaluationClient", "WritingEvaluationNormalizer", "normalize + validate", "call"),
                    ("WritingEvaluationCacheService", "Database", "persist status/score/evidence", "call"),
                ],
            ),
            uc(
                "UC-WRT-02", "Tái sử dụng hoặc yêu cầu chấm lại Writing", "CURRENT",
                ["Người học", "Người đánh giá"], ["Nhà cung cấp AI"],
                "Tái sử dụng evaluation cùng identity khi hợp lệ và chỉ chấm lại qua action có quyền/audit, không chấm lại khi mở trang kết quả.",
                ["Evaluation identity xác định được.", "Actor có quyền re-evaluate nếu yêu cầu chấm lại."],
                ["Kết quả cùng identity được reuse hoặc evaluation mới có provenance.", "GET result vẫn read-only."],
                ["Không chấm lại âm thầm.", "Không ghi đè history mà mất audit."],
                [
                    "Hệ thống tính evaluation identity từ task, answer và rubric version.",
                    "Cache service tìm evaluation reusable.",
                    "Nếu READY và còn hợp lệ, presenter dùng lại kết quả.",
                    "Nếu reviewer yêu cầu chấm lại, controller kiểm tra quyền và lý do.",
                    "Service tạo evaluation attempt mới có liên kết lịch sử.",
                    "UI hiển thị trạng thái mới mà vẫn giữ provenance trước đó.",
                ],
                ["A1 - Identity thay đổi: không reuse kết quả cũ.", "A2 - Actor thiếu quyền: từ chối re-evaluate và giữ nguyên evaluation hiện tại."],
                ["BR-WRT-04: Result GET không gọi provider.", "BR-WRT-05: Re-evaluate là command có quyền và audit.", "BR-WRT-06: Identity gồm immutable prompt/answer/rubric contract."],
                ["MSG-WRT-03: Đang sử dụng kết quả chấm phù hợp hiện có.", "MSG-WRT-04: Bạn không có quyền yêu cầu chấm lại."],
                ["WritingEvaluationCacheService", "PracticeController.reEvaluateQuestion", "WritingEvaluationReusePolicy"],
                ["Learner/Reviewer", "Result View", "PracticeController", "WritingEvaluationCacheService", "WritingEvaluationClient", "Database"],
                [
                    ("Learner/Reviewer", "Result View", "open result or request re-evaluate", "call"),
                    ("Result View", "PracticeController", "GET result / POST re-evaluate", "call"),
                    ("PracticeController", "WritingEvaluationCacheService", "resolve identity + permission", "call"),
                    ("WritingEvaluationCacheService", "Database", "find reusable evaluation/history", "call"),
                    ("WritingEvaluationCacheService", "WritingEvaluationClient", "call only for approved command", "call"),
                    ("WritingEvaluationCacheService", "Database", "persist new provenance when needed", "call"),
                    ("PracticeController", "Result View", "render reuse/pending/result", "return"),
                ],
            ),
            uc(
                "UC-WRT-03", "Xem bằng chứng rubric và cải thiện bài viết", "PLANNED 13E",
                ["Người học"], ["Writing evidence presenter"],
                "Trình bày bốn lăng kính phân tích phù hợp tiếng Hàn cùng bài gốc, correction, upgraded rewrite và sample mà không đổi công thức điểm chính thức.",
                ["Attempt Writing thuộc người học và có result context.", "Evaluation status xác định được."],
                ["Bốn lăng kính có evidence cụ thể và màu ngữ nghĩa.", "Official task score/weight vẫn là nguồn điểm duy nhất."],
                ["Không double-count tiêu chí.", "Không hiển thị correction/sample không có provenance."],
                [
                    "Người học mở Result Detail của bài Writing.",
                    "Detail assembler nạp answer text và evaluation immutable của attempt.",
                    "Presenter tách task response, organization/cohesion, vocabulary/expression và grammar/accuracy.",
                    "Mỗi lens hiển thị descriptor, evidence và semantic scale có nhãn.",
                    "UI hiển thị bài gốc, correction, rewrite nâng cấp và sample theo availability.",
                    "Trang giữ trạng thái pending/failed/unavailable rõ ràng mà không tạo nội dung giả.",
                ],
                ["A1 - Chưa chấm xong: hiển thị pending và CTA recovery được phép.", "A2 - Chỉ có feedback legacy: compatibility reader hiển thị giới hạn và gắn debt Phase 15."],
                ["BR-WRT-07: Bốn lens là presentation, không thay ba tiêu chí TOPIK chính thức.", "BR-WRT-08: Semantic color luôn kèm nhãn chữ.", "BR-WRT-09: 13E không tạo pipeline AI thứ hai."],
                ["MSG-WRT-05: Chi tiết chấm đang được chuẩn bị.", "MSG-WRT-06: Chưa có bản sửa hoặc bài mẫu cho lần chấm này."],
                ["PracticeResultDetailAssembler (planned)", "WritingEvidencePresenter (planned)", "WritingFeedbackCompatibilityReader"],
                ["Learner", "Writing Detail View", "PracticeController", "PracticeResultDetailAssembler [13E]", "WritingEvidencePresenter [13E]", "Database"],
                [
                    ("Learner", "Writing Detail View", "open Writing detail", "call"),
                    ("Writing Detail View", "PracticeController", "GET result/detail", "call"),
                    ("PracticeController", "PracticeResultDetailAssembler [13E]", "assemble immutable evidence", "call"),
                    ("PracticeResultDetailAssembler [13E]", "Database", "read answer + evaluation provenance", "call"),
                    ("PracticeResultDetailAssembler [13E]", "WritingEvidencePresenter [13E]", "map four Korean lenses", "call"),
                    ("WritingEvidencePresenter [13E]", "Writing Detail View", "render evidence and availability states", "return"),
                ],
            ),
        ],
    },
    {
        "code": "SPK",
        "name": "Media Nói, phiên âm và chấm AI",
        "short": "Speaking Evaluation",
        "purpose": "Lưu recording riêng tư theo attempt, phiên âm, chấm theo sáu tiêu chí tiếng Hàn và trình bày tổng quan holistic.",
        "actors": ["Người học", "Người đánh giá"],
        "secondary": ["Kho audio riêng tư", "Nhà cung cấp phiên âm/AI"],
        "classes": [
            ("PracticeSpeakingMediaController", "Controller", "CURRENT", ["upload/activate/delete", "owner-scoped endpoints"]),
            ("PracticeSpeakingMediaService", "Service", "CURRENT", ["media lifecycle", "attempt ownership"]),
            ("SpeakingAudioPreparationService", "Service", "CURRENT", ["validate/inspect audio", "prepare private object"]),
            ("SpeakingTranscriptionMediaResolver", "Service", "CURRENT", ["resolve authorized recording", "transcription input"]),
            ("SpeakingEvaluationOrchestrator", "Service", "CURRENT", ["transcribe/evaluate sequence", "failure policy"]),
            ("SpeakingRuleEngine", "Domain service", "CURRENT", ["six Korean criteria", "evidence checks"]),
            ("SpeakingResultPresenter", "Presenter", "CURRENT", ["holistic overview", "no per-question overview"]),
            ("SpeakingEvidencePresenter", "Logical presenter", "PLANNED 13E", ["per-question recording/transcript", "structured evidence"]),
        ],
        "relations": [
            ("PracticeSpeakingMediaController", "PracticeSpeakingMediaService", "manages"),
            ("PracticeSpeakingMediaService", "SpeakingAudioPreparationService", "validates/stores"),
            ("SpeakingEvaluationOrchestrator", "SpeakingTranscriptionMediaResolver", "loads audio"),
            ("SpeakingEvaluationOrchestrator", "SpeakingRuleEngine", "validates feedback"),
            ("SpeakingEvaluationOrchestrator", "SpeakingResultPresenter", "feeds"),
            ("SpeakingResultPresenter", "SpeakingEvidencePresenter", "will share context"),
        ],
        "use_cases": [
            uc(
                "UC-SPK-01", "Quản lý recording riêng tư theo attempt", "CURRENT",
                ["Người học"], ["Kho audio riêng tư"],
                "Kiểm tra, lưu, kích hoạt, phát lại và dọn recording theo owner/attempt/question mà không công khai đường dẫn file.",
                ["Người học sở hữu attempt Speaking đang cho phép upload.", "File audio thuộc định dạng/kích thước cho phép."],
                ["Recording hợp lệ được lưu private và gắn identity.", "Chỉ owner/reviewer có quyền mới phát lại được."],
                ["File độc hại/sai định dạng không được kích hoạt.", "Recording orphan có cleanup task."],
                [
                    "Player gửi recording cùng attempt/question identity.",
                    "Controller xác thực owner và state attempt.",
                    "Preparation service kiểm tra content, duration và ffprobe metadata.",
                    "Storage lưu object private và trả opaque storage key.",
                    "Media service kích hoạt recording hiện hành trong transaction.",
                    "Playback endpoint phục vụ byte-range sau authorization.",
                ],
                ["A1 - Audio không hợp lệ: xóa file tạm và trả lỗi có category.", "A2 - Thay recording: media cũ chuyển cleanup state, không lộ hoặc xóa nhầm media đang active."],
                ["BR-SPK-01: Speaking media mặc định private.", "BR-SPK-02: Authorization kiểm tra owner/attempt/reviewer.", "BR-SPK-03: Playback hỗ trợ byte range đúng hợp đồng."],
                ["MSG-SPK-01: Bản ghi âm không hợp lệ.", "MSG-SPK-02: Bạn không có quyền nghe bản ghi này."],
                ["PracticeSpeakingMediaController", "PracticeSpeakingMediaService", "SpeakingAudioPreparationService"],
                ["Learner", "Speaking Player", "PracticeSpeakingMediaController", "PracticeSpeakingMediaService", "SpeakingAudioPreparationService", "Private Storage"],
                [
                    ("Learner", "Speaking Player", "record and upload", "call"),
                    ("Speaking Player", "PracticeSpeakingMediaController", "POST private recording", "call"),
                    ("PracticeSpeakingMediaController", "PracticeSpeakingMediaService", "authorize media action", "call"),
                    ("PracticeSpeakingMediaService", "SpeakingAudioPreparationService", "inspect and prepare audio", "call"),
                    ("SpeakingAudioPreparationService", "Private Storage", "store private object", "call"),
                    ("Private Storage", "PracticeSpeakingMediaService", "opaque storage key", "return"),
                    ("PracticeSpeakingMediaService", "Speaking Player", "active recording identity", "return"),
                ],
            ),
            uc(
                "UC-SPK-02", "Phiên âm và chấm bài Nói", "CURRENT",
                ["Người học"], ["Nhà cung cấp phiên âm/AI"],
                "Resolve recording đã ủy quyền, phiên âm và đánh giá theo sáu tiêu chí tiếng Hàn với status/provenance rõ ràng.",
                ["Attempt Speaking đã submit.", "Recording active hợp lệ hoặc transcript có provenance được phép."],
                ["Transcript/evaluation hợp lệ được lưu.", "Holistic score và sáu rubric criteria có evidence."],
                ["Không chấm audio chưa được resolver cấp quyền.", "Không fabricate score khi provider lỗi."],
                [
                    "Submit service yêu cầu đánh giá Speaking.",
                    "Orchestrator xác định evidence source và evaluation identity.",
                    "Media resolver nạp recording private đã authorize.",
                    "Transcription client tạo transcript có provenance.",
                    "Evaluation client chấm transcript/evidence theo prompt rules.",
                    "Normalizer/rule engine kiểm tra sáu tiêu chí rồi lưu result/status.",
                ],
                ["A1 - Không có media hợp lệ: lưu unavailable/failure có lý do, không giả transcript.", "A2 - Transcript-only fallback: gắn evidence limitation; live multimodal audio scoring vẫn NO-GO."],
                ["BR-SPK-04: Media resolver là cửa duy nhất tới private recording cho AI.", "BR-SPK-05: Sáu tiêu chí phải khớp rubric tiếng Hàn.", "BR-SPK-06: Provider failure không fabricate score."],
                ["MSG-SPK-03: Bài nói đang được phiên âm và đánh giá.", "MSG-SPK-04: Chưa thể đánh giá bản ghi này."],
                ["SpeakingEvaluationOrchestrator", "SpeakingTranscriptionMediaResolver", "SpeakingRuleEngine"],
                ["Submit Service", "SpeakingEvaluationOrchestrator", "Media Resolver", "Transcription Provider", "Evaluation Provider", "SpeakingRuleEngine", "Database"],
                [
                    ("Submit Service", "SpeakingEvaluationOrchestrator", "evaluate submitted attempt", "call"),
                    ("SpeakingEvaluationOrchestrator", "Media Resolver", "resolve authorized recording", "call"),
                    ("Media Resolver", "Transcription Provider", "transcribe private audio", "call"),
                    ("Transcription Provider", "SpeakingEvaluationOrchestrator", "transcript + provenance", "return"),
                    ("SpeakingEvaluationOrchestrator", "Evaluation Provider", "evaluate Korean rubric", "call"),
                    ("Evaluation Provider", "SpeakingRuleEngine", "feedback/score candidate", "return"),
                    ("SpeakingRuleEngine", "Database", "persist validated status/evidence", "call"),
                ],
            ),
            uc(
                "UC-SPK-03", "Xem tổng quan holistic và evidence từng câu", "PLANNED 13E",
                ["Người học"], ["Speaking evidence presenter"],
                "Giữ overview là đánh giá toàn attempt, còn Result Detail cho phép kiểm tra recording/transcript/evidence từng câu mà không biến overview thành phân tích câu hỏi.",
                ["Attempt thuộc người học.", "Evaluation status và media authorization xác định được."],
                ["Overview hiển thị holistic sáu tiêu chí.", "Detail hiển thị recording/transcript theo question cùng evidence limitation."],
                ["Không lộ private media.", "Không hiển thị per-question score bịa khi evaluator chỉ trả holistic result."],
                [
                    "Người học mở Speaking Result hoặc Result Detail.",
                    "Overview presenter tổng hợp sáu tiêu chí trên toàn attempt.",
                    "Detail assembler nạp question delivery và media identities của attempt.",
                    "Media playback URL được tạo qua endpoint ủy quyền.",
                    "Evidence presenter ghép transcript/evidence theo câu và feedback holistic riêng.",
                    "UI hiển thị limitation khi chỉ có transcript hoặc thiếu recording.",
                ],
                ["A1 - Recording đã unavailable: giữ transcript/provenance nếu hợp lệ và hiển thị trạng thái media.", "A2 - Feedback legacy ESSAY/mixed: compatibility reader hiển thị giới hạn, dọn ở Phase 15."],
                ["BR-SPK-07: Overview không phân tích từng câu.", "BR-SPK-08: Detail không suy diễn per-question score từ holistic score.", "BR-SPK-09: Playback luôn qua authorization."],
                ["MSG-SPK-05: Bản ghi không còn khả dụng.", "MSG-SPK-06: Đánh giá hiện chỉ dựa trên transcript."],
                ["PracticeResultDetailAssembler (planned)", "SpeakingEvidencePresenter (planned)", "PracticeSpeakingMediaPlaybackService"],
                ["Learner", "Speaking Detail View", "PracticeController", "PracticeResultDetailAssembler [13E]", "SpeakingEvidencePresenter [13E]", "Private Media Endpoint"],
                [
                    ("Learner", "Speaking Detail View", "open Speaking detail", "call"),
                    ("Speaking Detail View", "PracticeController", "GET result/detail", "call"),
                    ("PracticeController", "PracticeResultDetailAssembler [13E]", "load holistic + question evidence", "call"),
                    ("PracticeResultDetailAssembler [13E]", "SpeakingEvidencePresenter [13E]", "map transcript/media/evidence", "call"),
                    ("SpeakingEvidencePresenter [13E]", "Private Media Endpoint", "create authorized playback references", "call"),
                    ("SpeakingEvidencePresenter [13E]", "Speaking Detail View", "render holistic and per-question evidence", "return"),
                ],
            ),
        ],
    },
    {
        "code": "RES",
        "name": "Kết quả tổng quan và Result Detail theo evidence",
        "short": "Result Overview & Detail",
        "purpose": "Dùng một result envelope/assembler, presenter theo kỹ năng và detail evidence tách learner answer, official key, teacher explanation và AI artifact.",
        "actors": ["Người học", "Người đánh giá"],
        "secondary": ["Explanation read service", "Kho media riêng tư"],
        "classes": [
            ("PracticeResultAssembler", "Assembler", "CURRENT", ["canonical result envelope", "select skill presenter"]),
            ("ObjectiveResultPresenter", "Presenter", "CURRENT", ["Reading/Listening overview", "question-type accuracy"]),
            ("WritingResultPresenter", "Presenter", "CURRENT", ["Korean Writing overview", "four analysis lenses"]),
            ("SpeakingResultPresenter", "Presenter", "CURRENT", ["holistic Speaking overview", "six criteria"]),
            ("PracticeResultContext", "Context", "CURRENT", ["attempt/version/evaluation data", "authorization-safe input"]),
            ("PracticeResultDetailAssembler", "Logical assembler", "PLANNED 13E", ["evidence layers", "availability states"]),
            ("ObjectiveEvidencePresenter", "Logical presenter", "PLANNED 13E", ["learner/key/teacher/AI separation", "anchors"]),
            ("Writing/SpeakingEvidencePresenter", "Logical presenters", "PLANNED 13E", ["skill-native detail", "no second AI pipeline"]),
        ],
        "relations": [
            ("PracticeResultAssembler", "PracticeResultContext", "consumes"),
            ("PracticeResultAssembler", "ObjectiveResultPresenter", "delegates"),
            ("PracticeResultAssembler", "WritingResultPresenter", "delegates"),
            ("PracticeResultAssembler", "SpeakingResultPresenter", "delegates"),
            ("PracticeResultContext", "PracticeResultDetailAssembler", "will feed"),
            ("PracticeResultDetailAssembler", "ObjectiveEvidencePresenter", "will delegate"),
            ("PracticeResultDetailAssembler", "Writing/SpeakingEvidencePresenter", "will delegate"),
        ],
        "use_cases": [
            uc(
                "UC-RES-01", "Xem Result Overview theo kỹ năng", "CURRENT",
                ["Người học"], ["Result presenters"],
                "Hiển thị shared header/envelope nhưng chọn đúng presenter: objective cho Đọc/Nghe, Korean-native cho Viết và holistic cho Nói.",
                ["Attempt ở trạng thái có result.", "Người học sở hữu attempt hoặc reviewer được phép."],
                ["Một overview phù hợp kỹ năng được hiển thị.", "Pending/failure status được giữ nguyên nghĩa."],
                ["Không dùng một bảng generic cho mọi kỹ năng.", "Không gọi AI khi mở GET."],
                [
                    "Người học mở route /result.",
                    "Controller xác thực quyền với attempt/version.",
                    "Result assembler tạo canonical context/envelope.",
                    "Assembler chọn objective, Writing hoặc Speaking presenter.",
                    "Presenter xây metrics/criteria phù hợp evidence hiện có.",
                    "Template shared frame hiển thị overview và CTA detail.",
                ],
                ["A1 - Evaluation pending/failed: hiển thị status card thay vì score giả.", "A2 - Skill không hỗ trợ: fail closed với unavailable state có diagnostic an toàn."],
                ["BR-RES-01: Chỉ một top-level result assembler.", "BR-RES-02: Exactly three wired overview presenter families.", "BR-RES-03: Speaking overview là holistic."],
                ["MSG-RES-01: Kết quả đang được hoàn thiện.", "MSG-RES-02: Chưa thể hiển thị kết quả này."],
                ["PracticeResultAssembler", "ObjectiveResultPresenter", "WritingResultPresenter", "SpeakingResultPresenter"],
                ["Learner", "Result View", "PracticeController", "PracticeResultAssembler", "Skill Result Presenter", "Database"],
                [
                    ("Learner", "Result View", "open overview", "call"),
                    ("Result View", "PracticeController", "GET /result", "call"),
                    ("PracticeController", "PracticeResultAssembler", "assemble canonical context", "call"),
                    ("PracticeResultAssembler", "Database", "read attempt/version/evaluation", "call"),
                    ("PracticeResultAssembler", "Skill Result Presenter", "present by skill family", "call"),
                    ("Skill Result Presenter", "Result View", "render overview/status", "return"),
                ],
            ),
            uc(
                "UC-RES-02", "Xem Result Detail tách lớp evidence", "PLANNED 13E",
                ["Người học"], ["Explanation read service", "Kho media riêng tư"],
                "Hiển thị từng lớp evidence đúng provenance: learner answer, official key, teacher explanation, immutable AI artifact và evidence riêng của Writing/Speaking.",
                ["Attempt/version ownership hợp lệ.", "Result detail route không tạo provider work."],
                ["Mỗi evidence layer có nhãn nguồn/trạng thái.", "Các kỹ năng dùng presenter detail phù hợp."],
                ["Không trộn official key với learner answer.", "Không fallback sang nội dung live/current khác version."],
                [
                    "Người học chọn Xem chi tiết.",
                    "Detail assembler nạp immutable question delivery của attempt.",
                    "Objective presenter tách answer/key/teacher/AI explanation và anchors.",
                    "Writing presenter tách submitted text/rubric/correction/rewrite/sample.",
                    "Speaking presenter giữ holistic feedback và ghép recording/transcript từng câu.",
                    "Template hiển thị availability/status mà không gọi pipeline AI mới.",
                ],
                ["A1 - Evidence chưa sẵn sàng: hiển thị pending/failed/unavailable theo loại.", "A2 - Historical compatibility data: đọc qua compatibility reader và ghi debt Phase 15, không che giấu provenance."],
                ["BR-RES-04: 13E là presentation layer trên 13D lifecycle.", "BR-RES-05: Evidence luôn gắn immutable version/provenance.", "BR-RES-06: GET detail không gọi provider."],
                ["MSG-RES-03: Evidence này đang được chuẩn bị.", "MSG-RES-04: Evidence không khả dụng cho phiên bản đã làm."],
                ["PracticeResultDetailAssembler (planned)", "ObjectiveEvidencePresenter (planned)", "WritingEvidencePresenter (planned)", "SpeakingEvidencePresenter (planned)"],
                ["Learner", "Result Detail View", "PracticeController", "PracticeResultDetailAssembler [13E]", "Skill Evidence Presenter [13E]", "Read Services", "Database"],
                [
                    ("Learner", "Result Detail View", "open evidence detail", "call"),
                    ("Result Detail View", "PracticeController", "GET /result/detail", "call"),
                    ("PracticeController", "PracticeResultDetailAssembler [13E]", "load immutable result context", "call"),
                    ("PracticeResultDetailAssembler [13E]", "Read Services", "read explanation/evaluation/media metadata", "call"),
                    ("Read Services", "Database", "fetch version-bound evidence", "call"),
                    ("PracticeResultDetailAssembler [13E]", "Skill Evidence Presenter [13E]", "map separated layers", "call"),
                    ("Skill Evidence Presenter [13E]", "Result Detail View", "render evidence/status", "return"),
                ],
            ),
            uc(
                "UC-RES-03", "Xử lý trạng thái pending, failed và unavailable", "CURRENT",
                ["Người học", "Người đánh giá"], ["Retry services"],
                "Biểu diễn trạng thái vận hành trung thực và chỉ cung cấp retry/re-evaluate cho actor có quyền qua command riêng.",
                ["Result context xác định trạng thái cho từng evaluation/artifact."],
                ["Learner thấy thông báo an toàn và nhất quán.", "Reviewer thấy action recovery được phép."],
                ["Không lộ provider payload nội bộ.", "Không biến GET thành retry command."],
                [
                    "Assembler đọc status của objective explanation, Writing và Speaking evaluation.",
                    "Presenter ánh xạ status thành learner-safe view model.",
                    "Learner thấy pending/failed/unavailable và dữ liệu vẫn khả dụng.",
                    "Reviewer có quyền có thể chọn retry/re-evaluate riêng.",
                    "Controller xác thực quyền và gửi command tới retry service.",
                    "UI tải lại trạng thái mới mà giữ audit/provenance.",
                ],
                ["A1 - Failure terminal: ẩn retry với learner và hiển thị đường liên hệ/report phù hợp.", "A2 - Retry đang chạy: command lặp là idempotent và không tạo task duplicate."],
                ["BR-RES-07: GET chỉ đọc.", "BR-RES-08: Recovery action có quyền và audit.", "BR-RES-09: Learner message không chứa provider/internal stack details."],
                ["MSG-RES-05: Kết quả đang xử lý.", "MSG-RES-06: Một phần đánh giá chưa khả dụng."],
                ["QuestionExplanationRetryService", "PracticeController.reEvaluate", "PracticeResultAssembler"],
                ["Learner/Reviewer", "Result View", "PracticeController", "PracticeResultAssembler", "Retry/Re-evaluate Service", "Database"],
                [
                    ("Learner/Reviewer", "Result View", "view state / choose recovery", "call"),
                    ("Result View", "PracticeController", "GET status or POST command", "call"),
                    ("PracticeController", "PracticeResultAssembler", "map learner-safe status", "call"),
                    ("PracticeController", "Retry/Re-evaluate Service", "authorize recovery command", "call"),
                    ("Retry/Re-evaluate Service", "Database", "transition durable status", "call"),
                    ("PracticeController", "Result View", "render updated state", "return"),
                ],
            ),
        ],
    },
    {
        "code": "PRG",
        "name": "Tiến độ học tập và phục hồi vận hành",
        "short": "Progress & Recovery",
        "purpose": "Hiển thị aggregate thật theo kỹ năng/thời gian, drill-down tới attempt và trạng thái recovery mà không bịa phần trăm trang trí.",
        "actors": ["Người học", "Người vận hành"],
        "secondary": ["Result services"],
        "classes": [
            ("PracticeController", "Controller", "CURRENT", ["/practice/progress route", "access/session handling"]),
            ("PracticeService", "Facade", "CURRENT", ["getProgressPageData", "bounded attempt summary"]),
            ("PracticeAttemptRepository", "Repository", "CURRENT", ["attempt history query", "status/date ordering"]),
            ("PracticeResultAssembler", "Assembler", "CURRENT", ["attempt result deep link", "skill status context"]),
            ("PracticeProgressQueryService", "Logical service", "PLANNED 13F", ["filters and real aggregates", "sample/recency/confidence"]),
            ("PracticeProgressAggregateRepository", "Logical repository", "PLANNED 13F", ["bounded aggregate queries", "partial credit support"]),
            ("PracticeProgressPresenter", "Logical presenter", "PLANNED 13F", ["charts/tables/empty states", "no fake metrics"]),
            ("PracticeRecoveryPresenter", "Logical presenter", "PLANNED 13F", ["pending/failure/retry UX", "authorized actions"]),
        ],
        "relations": [
            ("PracticeController", "PracticeService", "currently queries"),
            ("PracticeService", "PracticeAttemptRepository", "reads"),
            ("PracticeController", "PracticeProgressQueryService", "will query"),
            ("PracticeProgressQueryService", "PracticeProgressAggregateRepository", "will aggregate"),
            ("PracticeProgressQueryService", "PracticeProgressPresenter", "will feed"),
            ("PracticeProgressPresenter", "PracticeRecoveryPresenter", "will compose"),
            ("PracticeProgressPresenter", "PracticeResultAssembler", "deep-links"),
        ],
        "use_cases": [
            uc(
                "UC-PRG-01", "Xem lịch sử và aggregate thật", "CURRENT",
                ["Người học"], ["Result services"],
                "Hiển thị lịch sử attempt có giới hạn và aggregate chỉ từ dữ liệu thật; nền 13F sẽ bổ sung sample size, recency và confidence.",
                ["Người học đã đăng nhập.", "Attempt history có thể truy vấn theo owner."],
                ["Trang progress hiển thị lịch sử thật hoặc empty state.", "Không dùng seed giả trong production path."],
                ["Query lỗi trả trang recovery, không HTTP 500 thô."],
                [
                    "Người học mở /practice/progress.",
                    "Controller xác thực phiên và learner identity.",
                    "Service truy vấn attempt history có giới hạn.",
                    "Hệ thống tính các tổng hợp hiện có từ score/status thật.",
                    "Presenter hiển thị theo kỹ năng và thời gian khi đủ dữ liệu.",
                    "Người học có thể mở Result của một attempt cụ thể.",
                ],
                ["A1 - Không có attempt: hiển thị empty state và CTA về catalog.", "A2 - Query lỗi: hiển thị recovery state có request id thay vì response 500 không xử lý."],
                ["BR-PRG-01: Chỉ aggregate dữ liệu thật.", "BR-PRG-02: Query phải bounded.", "BR-PRG-03: Không so sánh kỹ năng khi thang điểm không tương thích."],
                ["MSG-PRG-01: Bạn chưa có kết quả luyện tập.", "MSG-PRG-02: Chưa thể tải tiến độ; vui lòng thử lại."],
                ["PracticeController.progress", "PracticeService.getProgressPageData", "PracticeAttemptRepository"],
                ["Learner", "Progress View", "PracticeController", "PracticeService", "PracticeAttemptRepository", "Database"],
                [
                    ("Learner", "Progress View", "open progress", "call"),
                    ("Progress View", "PracticeController", "GET /practice/progress", "call"),
                    ("PracticeController", "PracticeService", "getProgressPageData(owner)", "call"),
                    ("PracticeService", "PracticeAttemptRepository", "query bounded history", "call"),
                    ("PracticeAttemptRepository", "Database", "fetch attempt rows", "call"),
                    ("Database", "PracticeService", "history + scores/status", "return"),
                    ("PracticeController", "Progress View", "render aggregates/empty/recovery", "return"),
                ],
            ),
            uc(
                "UC-PRG-02", "Lọc và drill-down tiến độ", "PLANNED 13F",
                ["Người học"], ["Result services"],
                "Lọc lịch sử theo kỹ năng, khoảng thời gian, mode và trạng thái; hiển thị sample size/recency/confidence rồi deep-link tới result/detail.",
                ["13F progress query boundary được triển khai.", "Người học có ít nhất một attempt hoặc nhận empty state."],
                ["Bộ lọc ổn định trong URL.", "Aggregate và drill-down dùng cùng tập dữ liệu đã lọc."],
                ["Không hiển thị xu hướng có vẻ chính xác khi sample quá nhỏ."],
                [
                    "Người học chọn skill, date range, mode hoặc status.",
                    "Controller chuẩn hóa filter và giới hạn date range.",
                    "Progress query service chạy aggregate và history query cùng contract.",
                    "Service tính sample size, recency và confidence indicator.",
                    "Presenter hiển thị chart/table có nhãn và empty/low-confidence state.",
                    "Người học chọn một attempt để mở result hoặc evidence detail.",
                ],
                ["A1 - Filter không hợp lệ: dùng default an toàn và thông báo ngắn.", "A2 - Sample nhỏ: hiển thị dữ liệu thô với nhãn hạn chế, không suy diễn xu hướng."],
                ["BR-PRG-04: Aggregate/history phải dùng cùng filter contract.", "BR-PRG-05: Confidence luôn kèm sample size/recency.", "BR-PRG-06: Deep link giữ attempt identity, không dùng current version."],
                ["MSG-PRG-03: Chưa đủ dữ liệu để xác định xu hướng.", "MSG-PRG-04: Không có bài làm trong khoảng đã chọn."],
                ["PracticeProgressQueryService (planned)", "PracticeProgressPresenter (planned)", "PracticeResultAssembler"],
                ["Learner", "Progress View", "PracticeController", "PracticeProgressQueryService [13F]", "Aggregate Repository [13F]", "Result Routes"],
                [
                    ("Learner", "Progress View", "apply filters", "call"),
                    ("Progress View", "PracticeController", "GET progress?filters", "call"),
                    ("PracticeController", "PracticeProgressQueryService [13F]", "normalize + query", "call"),
                    ("PracticeProgressQueryService [13F]", "Aggregate Repository [13F]", "aggregate + history", "call"),
                    ("Aggregate Repository [13F]", "PracticeProgressQueryService [13F]", "rows + sample metadata", "return"),
                    ("PracticeProgressQueryService [13F]", "Progress View", "render metrics/confidence", "return"),
                    ("Progress View", "Result Routes", "open selected attempt", "call"),
                ],
            ),
            uc(
                "UC-PRG-03", "Phục hồi evaluation/explanation từ progress", "PLANNED 13F",
                ["Người học", "Người vận hành"], ["Retry/re-evaluate services"],
                "Đưa trạng thái pending/failed/unavailable vào progress và cho phép actor phù hợp đi tới recovery command hiện hữu, không tạo pipeline mới.",
                ["Progress data có operational status.", "Actor có quyền tương ứng với recovery action."],
                ["Learner thấy trạng thái an toàn và đường tới result.", "Operator có thể retry qua service 13D/AI hiện hữu."],
                ["Không lộ internal provider errors.", "Không retry tự động khi render progress."],
                [
                    "Progress query nạp trạng thái explanation/evaluation cùng attempt summary.",
                    "Recovery presenter ánh xạ thành learner-safe status.",
                    "Người học mở result để xem dữ liệu hiện có.",
                    "Operator/reviewer chọn retry hoặc re-evaluate khi có quyền.",
                    "Controller gọi service 13D/Writing/Speaking hiện hữu bằng command riêng.",
                    "Progress tải lại trạng thái mới và giữ lịch sử/audit.",
                ],
                ["A1 - Failure terminal: chỉ hiển thị report/support path.", "A2 - Command trùng: idempotently trả task/evaluation hiện tại."],
                ["BR-PRG-07: 13F tái sử dụng retry/re-evaluate hiện hữu.", "BR-PRG-08: GET progress không gây side effect.", "BR-PRG-09: Learner-safe status tách khỏi operator diagnostics."],
                ["MSG-PRG-05: Một số đánh giá đang chờ xử lý.", "MSG-PRG-06: Không thể thử lại trạng thái này."],
                ["PracticeRecoveryPresenter (planned)", "QuestionExplanationRetryService", "PracticeController.reEvaluate"],
                ["Learner/Operator", "Progress View", "PracticeController", "PracticeRecoveryPresenter [13F]", "Existing Retry Services", "Database"],
                [
                    ("Learner/Operator", "Progress View", "view state / choose recovery", "call"),
                    ("Progress View", "PracticeController", "GET progress or POST recovery", "call"),
                    ("PracticeController", "PracticeRecoveryPresenter [13F]", "map safe operational state", "call"),
                    ("PracticeController", "Existing Retry Services", "authorize explicit command", "call"),
                    ("Existing Retry Services", "Database", "persist durable transition", "call"),
                    ("PracticeController", "Progress View", "render refreshed state", "return"),
                ],
            ),
        ],
    },
]


# The ten capabilities are intentionally declared in product-flow order. The
# remaining AUT/XLS/PDF/PLY/RLE/WRT/SPK/RES/PRG entries above are data, not
# executable architecture claims. This assertion guards accidental omissions.
EXPECTED_CODES = ["CAT", "AUT", "XLS", "PDF", "PLY", "RLE", "WRT", "SPK", "RES", "PRG"]


def build_docx_model() -> list[dict[str, Any]]:
    """Overlay reader-facing English copy without changing Draw.io data."""

    model = deepcopy(CAPABILITIES)
    use_case_ids = {item["id"] for cap in model for item in cap["use_cases"]}
    if set(USE_CASE_COPY) != use_case_ids:
        missing = sorted(use_case_ids - set(USE_CASE_COPY))
        extra = sorted(set(USE_CASE_COPY) - use_case_ids)
        raise ValueError(f"English Use Case copy mismatch; missing={missing}, extra={extra}")
    if set(CAPABILITY_COPY) != set(EXPECTED_CODES):
        raise ValueError("English capability copy must cover all ten capabilities")

    actor_names = {
        actor
        for cap in model
        for actor in cap["actors"] + cap["secondary"]
    } | {
        actor
        for cap in model
        for item in cap["use_cases"]
        for actor in item["primary"] + item["secondary"]
    }
    unknown_actors = sorted(actor_names - set(ACTOR_NAMES))
    if unknown_actors:
        raise ValueError(f"Missing English actor names: {unknown_actors}")

    for cap in model:
        cap_copy = CAPABILITY_COPY[cap["code"]]
        cap["name"] = cap_copy["name"]
        cap["purpose"] = cap_copy["purpose"]
        cap["actors"] = [ACTOR_NAMES[name] for name in cap["actors"]]
        cap["secondary"] = [ACTOR_NAMES[name] for name in cap["secondary"]]
        for item in cap["use_cases"]:
            item_copy = USE_CASE_COPY[item["id"]]
            for field in (
                "title",
                "description",
                "preconditions",
                "success",
                "failure",
                "steps",
                "alternatives",
            ):
                item[field] = deepcopy(item_copy[field])
            item["primary"] = [ACTOR_NAMES[name] for name in item["primary"]]
            item["secondary"] = [ACTOR_NAMES[name] for name in item["secondary"]]
            item["rules"] = [
                (rule_id, BUSINESS_RULES[rule_id])
                for rule_id in item_copy["rule_ids"]
            ]
            item["messages"] = [
                (message_id, SYSTEM_MESSAGES[message_id])
                for message_id in item_copy["message_ids"]
            ]
    return model


DOCX_CAPABILITIES = build_docx_model()


def build_docx_groups() -> list[dict[str, Any]]:
    """Group the ten technical modules into four reader-facing business areas."""

    capability_by_code = {cap["code"]: cap for cap in DOCX_CAPABILITIES}
    covered_codes = [code for group in DOCUMENT_GROUPS for code in group["capabilities"]]
    if sorted(covered_codes) != sorted(EXPECTED_CODES) or len(covered_codes) != len(set(covered_codes)):
        raise ValueError("Document groups must cover every technical capability exactly once")

    groups: list[dict[str, Any]] = []
    for group_copy in DOCUMENT_GROUPS:
        group = deepcopy(group_copy)
        group["modules"] = []
        group["use_cases"] = []
        for code in group_copy["capabilities"]:
            capability = capability_by_code[code]
            group["modules"].append({"code": code, "name": capability["name"]})
            for source_item in capability["use_cases"]:
                item = deepcopy(source_item)
                item["module_code"] = code
                item["module_name"] = capability["name"]
                group["use_cases"].append(item)
        groups.append(group)
    return groups


DOCX_GROUPS = build_docx_groups()


def validate_model() -> None:
    codes = [cap["code"] for cap in CAPABILITIES]
    if codes != EXPECTED_CODES:
        raise ValueError(f"Capability order mismatch: {codes}")
    use_case_ids: list[str] = []
    for cap in CAPABILITIES:
        if len(cap["use_cases"]) != 3:
            raise ValueError(f"{cap['code']} must own exactly three Use Cases")
        if len(cap["classes"]) < 6:
            raise ValueError(f"{cap['code']} needs a meaningful class boundary")
        for item in cap["use_cases"]:
            use_case_ids.append(item["id"])
            if item["status"] not in STATUS_STYLE:
                raise ValueError(f"Unknown status on {item['id']}: {item['status']}")
            participant_set = set(item["participants"])
            for source, target, _label, _kind in item["sequence"]:
                if source not in participant_set or target not in participant_set:
                    raise ValueError(f"Unknown participant in {item['id']}: {source} -> {target}")
    if len(use_case_ids) != 30 or len(set(use_case_ids)) != 30:
        raise ValueError("The architecture baseline must contain 30 unique Use Cases")
    if len(DOCUMENT_GROUPS) != 4:
        raise ValueError("The DOCX must expose exactly four functional groups")
    expected_rule_ids = [f"BR-{index:02d}" for index in range(1, len(BUSINESS_RULES) + 1)]
    if list(BUSINESS_RULES) != expected_rule_ids:
        raise ValueError("Business Rules must use a continuous global sequence")
    assigned_rule_ids = [
        rule_id
        for group in DOCX_GROUPS
        for item in group["use_cases"]
        for rule_id, _statement in item["rules"]
    ]
    if set(assigned_rule_ids) != set(expected_rule_ids):
        raise ValueError("Every Business Rule must be referenced by at least one Use Case")
    assigned_message_ids = [
        message_id
        for group in DOCX_GROUPS
        for item in group["use_cases"]
        for message_id, _statement in item["messages"]
    ]
    if set(assigned_message_ids) != set(SYSTEM_MESSAGES):
        raise ValueError("Every System Message must be referenced by at least one Use Case")
    for group in DOCX_GROUPS:
        for item in group["use_cases"]:
            if not 1 <= len(item["rules"]) <= 6:
                raise ValueError(f"{item['id']} must reference an appropriate variable number of Business Rules")
            if not 1 <= len(item["messages"]) <= 4:
                raise ValueError(f"{item['id']} must reference an appropriate variable number of System Messages")
    reader_copy = []
    for group in DOCX_GROUPS:
        reader_copy.extend([group["name"], group["purpose"]])
        for item in group["use_cases"]:
            reader_copy.extend(item["primary"] + item["secondary"])
            reader_copy.extend(
                [item["title"], item["description"]]
                + item["preconditions"]
                + item["success"]
                + item["failure"]
                + item["steps"]
                + item["alternatives"]
                + [statement for _rule_id, statement in item["rules"]]
                + [statement for _message_id, statement in item["messages"]]
            )
    if any("Learner" in text or "learner" in text for text in reader_copy):
        raise ValueError("Reader-facing DOCX copy must use Student instead of Learner")


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(run, *, size: float = 11, color: str = "0B2545", bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent: int = 120) -> None:
    if sum(widths) != 9360:
        raise ValueError(f"Table widths must total 9360 DXA: {widths}")
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=COLORS["muted"])
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    text_node = OxmlElement("w:t")
    text_node.text = "1"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_separate, text_node, fld_char_end])


class NumberingFactory:
    """Create real, restartable single-level Word numbering definitions."""

    def __init__(self, doc: Document):
        self.root = doc.part.numbering_part.element

    def create(self, kind: str) -> int:
        if kind not in {"decimal", "bullet"}:
            raise ValueError(f"Unsupported numbering kind: {kind}")
        abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in self.root.findall(qn("w:abstractNum"))]
        num_ids = [int(node.get(qn("w:numId"))) for node in self.root.findall(qn("w:num"))]
        abstract_id = max(abstract_ids, default=-1) + 1
        num_id = max(num_ids, default=0) + 1

        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        level.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), kind)
        level.append(num_fmt)
        level_text = OxmlElement("w:lvlText")
        level_text.set(qn("w:val"), "%1." if kind == "decimal" else "•")
        level.append(level_text)
        level_jc = OxmlElement("w:lvlJc")
        level_jc.set(qn("w:val"), "left")
        level.append(level_jc)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "360")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        p_pr.append(ind)
        level.append(p_pr)
        if kind == "bullet":
            r_pr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), "Arial")
            fonts.set(qn("w:hAnsi"), "Arial")
            r_pr.append(fonts)
            level.append(r_pr)
        abstract.append(level)

        first_num = self.root.find(qn("w:num"))
        if first_num is None:
            self.root.append(abstract)
        else:
            self.root.insert(list(self.root).index(first_num), abstract)

        number = OxmlElement("w:num")
        number.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        number.append(abstract_ref)
        self.root.append(number)
        return num_id

    @staticmethod
    def apply(paragraph, num_id: int) -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is None:
            num_pr = OxmlElement("w:numPr")
            p_pr.append(num_pr)
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(num_id_node)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(COLORS["ink"])
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, COLORS["blue"], 16, 8),
        "Heading 2": (13, COLORS["blue"], 12, 6),
        "Heading 3": (12, COLORS["blue_dark"], 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(9)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(1)
        style.paragraph_format.line_spacing = 1.0

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run("KSH PRACTICE  |  USE CASE SPECIFICATIONS  |  PRE-13E")
    set_run_font(run, size=8.5, color=COLORS["muted"], bold=True)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    add_page_number(fp)


def add_title_paragraph(doc: Document, text: str, size: float, color: str, *, bold: bool = True, after: float = 6, align=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)
    return paragraph


def add_metadata_line(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    label_run = p.add_run(f"{label}: ")
    set_run_font(label_run, size=10.5, color=COLORS["ink"], bold=True)
    value_run = p.add_run(value)
    set_run_font(value_run, size=10.5, color=COLORS["ink"])


def add_cell_plain(cell, text: str, *, bold: bool = False, color: str = "0B2545", size: float = 9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)


def add_cell_list(cell, items: list[str], numbering: NumberingFactory, *, numbered: bool = False, heading: str | None = None) -> None:
    cell.text = ""
    if heading:
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(heading)
        set_run_font(run, size=9, color=COLORS["ink"], bold=True)
    else:
        cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
    num_id = numbering.create("decimal" if numbered else "bullet")
    for item in items:
        p = cell.add_paragraph()
        NumberingFactory.apply(p, num_id)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(item)
        set_run_font(run, size=8.7, color=COLORS["ink"])


def add_postconditions(cell, success_items: list[str], failure_items: list[str], numbering: NumberingFactory) -> None:
    cell.text = ""
    for heading, items, color in (("Success", success_items, COLORS["success"]), ("Failure-safe outcome", failure_items, COLORS["error"])):
        p = cell.add_paragraph() if cell.paragraphs[0].text else cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(heading)
        set_run_font(run, size=8.8, color=color, bold=True)
        num_id = numbering.create("bullet")
        for item in items:
            lp = cell.add_paragraph()
            NumberingFactory.apply(lp, num_id)
            lp.paragraph_format.space_after = Pt(1)
            lp.paragraph_format.line_spacing = 1.0
            item_run = lp.add_run(item)
            set_run_font(item_run, size=8.6, color=COLORS["ink"])


def add_business_rules(cell, items: list[tuple[str, str]]) -> None:
    """Render plain-language rule ids and statements as compact definition blocks."""

    cell.text = ""
    for index, (rule_id, statement) in enumerate(items):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.05
        id_run = paragraph.add_run(rule_id)
        set_run_font(id_run, size=8.7, color=COLORS["blue_dark"], bold=True)
        id_run.add_break()
        statement_run = paragraph.add_run(statement)
        set_run_font(statement_run, size=8.7, color=COLORS["ink"])


def add_system_messages(cell, items: list[tuple[str, str]]) -> None:
    """Render shared message ids and their Student-facing wording."""

    cell.text = ""
    for index, (message_id, statement) in enumerate(items):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.0
        id_run = paragraph.add_run(f"{message_id}: ")
        set_run_font(id_run, size=8.7, color=COLORS["blue_dark"], bold=True)
        statement_run = paragraph.add_run(statement)
        set_run_font(statement_run, size=8.7, color=COLORS["ink"])


def add_use_case_table(doc: Document, item: dict[str, Any], numbering: NumberingFactory) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_table_geometry(table, [1800, 2880, 1800, 2880])

    actor_row = table.rows[0]
    labels = ("Primary Actors", ", ".join(item["primary"]), "Secondary Actors", ", ".join(item["secondary"]) or "None")
    for index, value in enumerate(labels):
        is_label = index in (0, 2)
        add_cell_plain(actor_row.cells[index], value, bold=is_label, color=COLORS["blue_dark"] if is_label else COLORS["ink"], size=8.7)
        if is_label:
            set_cell_shading(actor_row.cells[index], COLORS["blue_light"])

    def add_row(label: str):
        row = table.add_row()
        label_cell = row.cells[0]
        detail_cell = row.cells[1].merge(row.cells[3])
        add_cell_plain(label_cell, label, bold=True, color=COLORS["blue_dark"], size=8.7)
        set_cell_shading(label_cell, COLORS["blue_light"])
        return detail_cell

    add_cell_plain(add_row("Description"), item["description"], size=8.8)
    add_cell_list(add_row("Preconditions"), item["preconditions"], numbering)
    add_postconditions(add_row("Postconditions"), item["success"], item["failure"], numbering)
    add_cell_list(add_row("Normal Sequence / Flow"), item["steps"], numbering, numbered=True)
    add_cell_list(add_row("Alternative Sequences / Flows"), item["alternatives"], numbering)
    add_business_rules(add_row("Business Rules"), item["rules"])
    add_system_messages(add_row("System Messages"), item["messages"])

    # Merged rows need explicit aggregate widths after Word has rewritten tcW.
    for row in table.rows[1:]:
        widths = [1800, 7560]
        for index, cell in enumerate(row.cells[:2]):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def add_document_bullets(doc: Document, items: list[str], *, size: float = 8.8) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(item)
        set_run_font(run, size=size, color=COLORS["ink"])


def add_jira_delivery_stream_page(doc: Document, stream: dict[str, Any]) -> None:
    doc.add_heading(
        f"B.{stream['order']} Task {stream['order']} - {stream['name']}",
        level=1,
    )

    summary = doc.add_table(rows=4, cols=2)
    summary.style = "Table Grid"
    set_table_geometry(summary, [2100, 7260])
    summary_rows = [
        ("Suggested Task", stream["summary"]),
        ("Traceability modules", stream["modules"]),
        ("Depends on", stream["depends_on"]),
        ("Blocks", stream["blocks"]),
    ]
    for row, (label, value) in zip(summary.rows, summary_rows):
        add_cell_plain(row.cells[0], label, bold=True, color=COLORS["blue_dark"], size=8.3)
        set_cell_shading(row.cells[0], COLORS["blue_light"])
        add_cell_plain(row.cells[1], value, size=8.3)

    doc.add_heading("Description", level=2)
    description = doc.add_paragraph(stream["description"])
    description.paragraph_format.space_after = Pt(4)
    for run in description.runs:
        set_run_font(run, size=9.0, color=COLORS["ink"])

    doc.add_heading("Acceptance boundary", level=2)
    add_document_bullets(doc, stream["acceptance"], size=8.4)

    doc.add_heading("Ordered Sub-tasks", level=2)
    subtasks = doc.add_table(rows=1, cols=3)
    subtasks.style = "Table Grid"
    set_table_geometry(subtasks, [600, 3480, 5280])
    for index, header in enumerate(("#", "Suggested Sub-task", "Deliverable / Done condition")):
        add_cell_plain(subtasks.rows[0].cells[index], header, bold=True, color=COLORS["ink"], size=7.8)
        set_cell_shading(subtasks.rows[0].cells[index], COLORS["blue_light"])
    set_repeat_table_header(subtasks.rows[0])
    for order, title, done in stream["subtasks"]:
        row = subtasks.add_row()
        add_cell_plain(row.cells[0], order, bold=True, color=COLORS["blue_dark"], size=7.4)
        add_cell_plain(row.cells[1], title, bold=True, size=7.4)
        add_cell_plain(row.cells[2], done, size=7.3)

    doc.add_heading("Linked Bug templates", level=2)
    bugs = doc.add_table(rows=1, cols=3)
    bugs.style = "Table Grid"
    set_table_geometry(bugs, [1080, 4680, 3600])
    for index, header in enumerate(("Risk", "Suggested Bug summary", "Placement / release effect")):
        add_cell_plain(bugs.rows[0].cells[index], header, bold=True, color=COLORS["ink"], size=7.7)
        set_cell_shading(bugs.rows[0].cells[index], COLORS["blue_light"])
    set_repeat_table_header(bugs.rows[0])
    for risk, title, placement in stream["bugs"]:
        row = bugs.add_row()
        add_cell_plain(row.cells[0], risk, bold=True, color=COLORS["error"], size=7.2)
        add_cell_plain(row.cells[1], title, bold=True, size=7.2)
        add_cell_plain(row.cells[2], placement, size=7.1)


def generate_docx() -> None:
    doc = Document()
    configure_document(doc)
    numbering = NumberingFactory(doc)
    doc.core_properties.title = "KSH Practice Use Case Specifications"
    doc.core_properties.subject = "Pre-13E four-area functional baseline"
    doc.core_properties.author = "KSH Engineering"
    doc.core_properties.keywords = "KSH, Practice, Use Case, Phase 13E, Architecture"

    # memo_masthead pattern, resolved with the standard_business_brief preset.
    doc.add_paragraph().paragraph_format.space_after = Pt(16)
    add_title_paragraph(doc, "KSH PRACTICE", 12, COLORS["blue"], after=12)
    add_title_paragraph(doc, "Use Case Specifications", 25, COLORS["ink"], after=4)
    add_title_paragraph(doc, "Four-area functional baseline before Phase 13E", 14, COLORS["muted"], bold=False, after=18)
    add_metadata_line(doc, "Branch", "feature/practice-reduce-scope")
    add_metadata_line(doc, "Status", "PRE_13E_ARCHITECTURE_BASELINE")
    add_metadata_line(doc, "Date", "18/07/2026")
    add_metadata_line(doc, "Scope", "Only /practice: 4 functional groups, 10 traceability modules and 30 formal Use Cases")
    add_metadata_line(doc, "Authority", "Current Spring Boot code plus approved Phase 13E-13H roadmap")
    doc.add_paragraph().paragraph_format.space_after = Pt(14)
    callout = doc.add_table(rows=1, cols=1)
    callout.style = "Table Grid"
    set_table_geometry(callout, [9360])
    set_cell_shading(callout.cell(0, 0), COLORS["panel"])
    add_cell_plain(
        callout.cell(0, 0),
        "Evidence convention: CURRENT means implemented code. PLANNED 13E/13F identifies an approved boundary that is not production code yet. Business Rules and System Messages use shared global identifiers and may be referenced by more than one Use Case.",
        size=10,
    )
    doc.add_page_break()

    doc.add_heading("1. Functional Group Map", level=1)
    p = doc.add_paragraph(
        "Practice is presented through four business-facing areas. The ten technical modules remain as traceability labels so the document stays compact without losing its connection to code, class diagrams or sequence diagrams."
    )
    p.paragraph_format.space_after = Pt(8)
    matrix = doc.add_table(rows=1, cols=4)
    matrix.style = "Table Grid"
    set_table_geometry(matrix, [720, 2500, 2600, 3540])
    headers = ["Code", "Functional area", "Internal modules", "Use Cases"]
    for i, header in enumerate(headers):
        add_cell_plain(matrix.rows[0].cells[i], header, bold=True, color=COLORS["ink"], size=9)
        set_cell_shading(matrix.rows[0].cells[i], COLORS["blue_light"])
    set_repeat_table_header(matrix.rows[0])
    for group in DOCX_GROUPS:
        row = matrix.add_row()
        module_labels = ", ".join(module["code"] for module in group["modules"])
        use_case_ids = ", ".join(item["id"] for item in group["use_cases"])
        values = [group["code"], group["name"], module_labels, use_case_ids]
        for i, value in enumerate(values):
            add_cell_plain(row.cells[i], value, bold=(i == 0), color=COLORS["blue_dark"] if i == 0 else COLORS["ink"], size=8.8)
    doc.add_heading("2. Shared Architecture Rules", level=1)
    rules = [
        "Attempt rendering and scoring always use the immutable version locked to the attempt.",
        "Reading and Listening result reads are read-only; provider work runs only through the durable preparation lifecycle.",
        "Student answers, official answers, teacher explanations and shared AI artifacts remain separate evidence layers.",
        "Writing or Speaking provider failure never creates a fabricated score or evidence item.",
        "Speaking media is private, and every upload or playback request requires authorization.",
        "Phase 13E adds evidence presentation; Phase 13F adds real progress aggregation and recovery behavior.",
    ]
    for rule in rules:
        p = doc.add_paragraph(rule, style="List Bullet")
        for run in p.runs:
            set_run_font(run, size=10, color=COLORS["ink"])
    doc.add_page_break()

    use_case_index = 0
    for group_index, group in enumerate(DOCX_GROUPS, start=1):
        for item in group["use_cases"]:
            use_case_index += 1
            kicker = doc.add_paragraph()
            kicker.paragraph_format.space_before = Pt(0)
            kicker.paragraph_format.space_after = Pt(2)
            kicker.paragraph_format.keep_with_next = True
            run = kicker.add_run(
                f"AREA {group_index:02d}/04  |  {group['code']} {group['name']}  |  MODULE {item['module_code']}  |  USE CASE {use_case_index:02d}/30"
            )
            set_run_font(run, size=8.5, color=COLORS["muted"], bold=True)
            add_title_paragraph(doc, f"{item['id']}  {item['title']}", 15, COLORS["ink"], after=3)
            status_p = doc.add_paragraph()
            status_p.paragraph_format.space_before = Pt(0)
            status_p.paragraph_format.space_after = Pt(3)
            status_p.paragraph_format.keep_with_next = True
            status_run = status_p.add_run(item["status"])
            fill, stroke = STATUS_STYLE[item["status"]]
            set_run_font(status_run, size=8.5, color=stroke, bold=True)
            source_run = status_p.add_run("   Code/plan trace: " + ", ".join(item["source"]))
            set_run_font(source_run, size=8.2, color=COLORS["muted"], italic=True)
            add_use_case_table(doc, item, numbering)
            if use_case_index < 30:
                doc.add_page_break()

    doc.add_page_break()
    doc.add_heading("Appendix A. Functional Group Inventory", level=1)
    doc.add_paragraph(
        "The DOCX uses four reader-facing areas while preserving the stable technical module and Use Case codes used by implementation and architecture artifacts. Business Rules and System Messages are shared catalogs: a single identifier can apply to several Use Cases."
    )
    appendix = doc.add_table(rows=1, cols=3)
    appendix.style = "Table Grid"
    set_table_geometry(appendix, [2700, 1800, 4860])
    for i, header in enumerate(("Functional area", "Internal modules", "Use Cases")):
        add_cell_plain(appendix.rows[0].cells[i], header, bold=True, size=9)
        set_cell_shading(appendix.rows[0].cells[i], COLORS["blue_light"])
    set_repeat_table_header(appendix.rows[0])
    for group in DOCX_GROUPS:
        row = appendix.add_row()
        values = [
            f"{group['code']} - {group['name']}",
            ", ".join(module["code"] for module in group["modules"]),
            ", ".join(item["id"] for item in group["use_cases"]),
        ]
        for i, value in enumerate(values):
            add_cell_plain(row.cells[i], value, bold=(i == 0), size=8.8)

    doc.add_page_break()
    doc.add_heading("Appendix B. Recommended Jira Delivery Hierarchy", level=1)
    intro = doc.add_paragraph(
        "This appendix defines a prospective implementation order for the Practice feature. It is a dependency-driven delivery plan, not a reconstruction of when work happened. Jira creation dates, sprint placement and work logs must remain truthful and must never be backdated from timestamps in Markdown files."
    )
    intro.paragraph_format.space_after = Pt(8)
    for run in intro.runs:
        set_run_font(run, size=9.5, color=COLORS["ink"])

    callout = doc.add_table(rows=1, cols=1)
    callout.style = "Table Grid"
    set_table_geometry(callout, [9360])
    set_cell_shading(callout.cell(0, 0), COLORS["panel"])
    add_cell_plain(
        callout.cell(0, 0),
        "Recommended hierarchy: one Practice Epic or project scope, four business-facing parent Tasks, ordered Sub-tasks inside each Task, and reproducible Bugs linked to the affected Task. The ten technical module codes remain Components, labels and traceability references; they are not ten competing delivery streams.",
        size=9.0,
    )

    doc.add_heading("Dependency order", level=2)
    dependency_table = doc.add_table(rows=1, cols=5)
    dependency_table.style = "Table Grid"
    set_table_geometry(dependency_table, [600, 800, 2700, 2560, 2700])
    for index, header in enumerate(("#", "Code", "Parent Task", "Depends on", "Why this order matters")):
        add_cell_plain(dependency_table.rows[0].cells[index], header, bold=True, size=8.0)
        set_cell_shading(dependency_table.rows[0].cells[index], COLORS["blue_light"])
    set_repeat_table_header(dependency_table.rows[0])
    dependency_reasons = {
        "MGT": "Published immutable content and version contracts are the input to an attempt.",
        "ATT": "A submitted, version-locked attempt is the authoritative input to result evidence.",
        "RSL": "Normalized authorized results are the authoritative input to progress aggregation.",
        "PRG": "Progress completes the read-side loop and may then guide a new authorized attempt.",
    }
    for stream in JIRA_DELIVERY_STREAMS:
        row = dependency_table.add_row()
        values = [
            str(stream["order"]),
            stream["code"],
            stream["name"],
            stream["depends_on"],
            dependency_reasons[stream["code"]],
        ]
        for index, value in enumerate(values):
            add_cell_plain(
                row.cells[index],
                value,
                bold=(index in (0, 1)),
                color=COLORS["blue_dark"] if index in (0, 1) else COLORS["ink"],
                size=7.7,
            )

    doc.add_heading("Issue hierarchy and placement rules", level=2)
    hierarchy = doc.add_table(rows=1, cols=3)
    hierarchy.style = "Table Grid"
    set_table_geometry(hierarchy, [1500, 2820, 5040])
    for index, header in enumerate(("Issue type", "Recommended role", "Placement rule")):
        add_cell_plain(hierarchy.rows[0].cells[index], header, bold=True, size=8.0)
        set_cell_shading(hierarchy.rows[0].cells[index], COLORS["blue_light"])
    set_repeat_table_header(hierarchy.rows[0])
    hierarchy_rows = [
        (
            "Epic / project scope",
            "Practice delivery umbrella",
            "Contains the four parent Tasks when the project uses Epics; do not mix unrelated Class Tests or Assignments into this scope.",
        ),
        (
            "Task",
            "One of the four delivery streams",
            "Owns description, acceptance boundary, dependency and the ordered implementation Sub-tasks shown below.",
        ),
        (
            "Sub-task",
            "One verifiable deliverable",
            "Use for requirements, diagrams, architecture, backend, frontend, tests and UAT. Keep each Done condition independently reviewable.",
        ),
        (
            "Bug",
            "One reproduced defect",
            "Create as a standalone Bug and link it to the affected Task with blocks or relates to. Do not create fictional Bugs or a catch-all fix-bugs Sub-task.",
        ),
    ]
    for issue_type, role, placement in hierarchy_rows:
        row = hierarchy.add_row()
        add_cell_plain(row.cells[0], issue_type, bold=True, color=COLORS["blue_dark"], size=7.8)
        add_cell_plain(row.cells[1], role, size=7.8)
        add_cell_plain(row.cells[2], placement, size=7.7)

    doc.add_heading("Bug evidence and release policy", level=2)
    add_document_bullets(
        doc,
        [
            "Create a Bug only after the defect is reproduced. Record steps, expected and actual behavior, evidence, affected Use Case and version, severity, and the regression test that proves the fix.",
            "Use blocks when the defect invalidates a contract, security boundary, version integrity, timing, scoring or downstream data. Use relates to for a contained non-blocking defect.",
            "Resolve contract and integrity Bugs before dependent-stream implementation; resolve integration and UI Bugs before UAT; resolve every release blocker before the parent Task is Done.",
            "The Bug summaries below are reusable risk templates, not claims that those defects currently exist.",
        ],
        size=8.5,
    )

    for stream in JIRA_DELIVERY_STREAMS:
        doc.add_page_break()
        add_jira_delivery_stream_page(doc, stream)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)


class DrawioPage:
    def __init__(self, parent: ET.Element, name: str, page_index: int):
        self.diagram = ET.SubElement(parent, "diagram", {"id": f"page-{page_index:02d}", "name": name})
        self.model = ET.SubElement(
            self.diagram,
            "mxGraphModel",
            {
                "dx": "1600",
                "dy": "1000",
                "grid": "1",
                "gridSize": "10",
                "guides": "1",
                "tooltips": "1",
                "connect": "1",
                "arrows": "1",
                "fold": "1",
                "page": "1",
                "pageScale": "1",
                "pageWidth": str(PAGE_WIDTH),
                "pageHeight": str(PAGE_HEIGHT),
                "math": "0",
                "shadow": "0",
            },
        )
        self.root = ET.SubElement(self.model, "root")
        ET.SubElement(self.root, "mxCell", {"id": "0"})
        ET.SubElement(self.root, "mxCell", {"id": "1", "parent": "0"})
        self.prefix = f"p{page_index:02d}"
        self.counter = 0

    def next_id(self) -> str:
        self.counter += 1
        return f"{self.prefix}-{self.counter:03d}"

    def vertex(self, value: str, x: float, y: float, width: float, height: float, style: str, *, parent: str = "1") -> str:
        cell_id = self.next_id()
        cell = ET.SubElement(self.root, "mxCell", {"id": cell_id, "value": value, "style": style, "vertex": "1", "parent": parent})
        ET.SubElement(cell, "mxGeometry", {"x": f"{x:.1f}", "y": f"{y:.1f}", "width": f"{width:.1f}", "height": f"{height:.1f}", "as": "geometry"})
        return cell_id

    def edge(self, source: str, target: str, label: str, style: str) -> str:
        cell_id = self.next_id()
        cell = ET.SubElement(self.root, "mxCell", {"id": cell_id, "value": label, "style": style, "edge": "1", "parent": "1", "source": source, "target": target})
        ET.SubElement(cell, "mxGeometry", {"relative": "1", "as": "geometry"})
        return cell_id

    def absolute_edge(self, x1: float, y1: float, x2: float, y2: float, label: str, style: str) -> str:
        cell_id = self.next_id()
        cell = ET.SubElement(self.root, "mxCell", {"id": cell_id, "value": label, "style": style, "edge": "1", "parent": "1"})
        geometry = ET.SubElement(cell, "mxGeometry", {"relative": "1", "as": "geometry"})
        ET.SubElement(geometry, "mxPoint", {"x": f"{x1:.1f}", "y": f"{y1:.1f}", "as": "sourcePoint"})
        ET.SubElement(geometry, "mxPoint", {"x": f"{x2:.1f}", "y": f"{y2:.1f}", "as": "targetPoint"})
        return cell_id


def drawio_text_style(size: int = 14, color: str = "0B2545", bold: bool = False, align: str = "left") -> str:
    return f"text;html=1;strokeColor=none;fillColor=none;align={align};verticalAlign=middle;whiteSpace=wrap;rounded=0;fontSize={size};fontColor=#{color};fontStyle={'1' if bold else '0'};spacing=4;"


def status_box_style(status: str, *, rounded: bool = True) -> str:
    fill, stroke = STATUS_STYLE[status]
    return (
        f"rounded={'1' if rounded else '0'};whiteSpace=wrap;html=1;fillColor=#{fill};strokeColor=#{stroke};"
        f"fontColor=#{COLORS['ink']};fontSize=12;align=left;verticalAlign=top;spacing=10;arcSize=8;"
    )


def add_page_header(page: DrawioPage, title: str, subtitle: str) -> None:
    page.vertex(title, 60, 28, 1480, 42, drawio_text_style(24, COLORS["ink"], True, "left"))
    page.vertex(subtitle, 60, 70, 1480, 30, drawio_text_style(12, COLORS["muted"], False, "left"))
    page.absolute_edge(60, 112, 1540, 112, "", f"endArrow=none;strokeColor=#{COLORS['line']};strokeWidth=1;")


def add_status_legend(page: DrawioPage, y: float = 930) -> None:
    x = 60
    for status in ("CURRENT", "PLANNED 13E", "PLANNED 13F", "DEFERRED 13H", "DEFERRED 15"):
        fill, stroke = STATUS_STYLE[status]
        page.vertex(status, x, y, 180, 32, f"rounded=1;whiteSpace=wrap;html=1;fillColor=#{fill};strokeColor=#{stroke};fontSize=10;fontStyle=1;fontColor=#{COLORS['ink']};")
        x += 200


def add_feature_decomposition(mxfile: ET.Element, index: int) -> None:
    page = DrawioPage(mxfile, "00 Feature Decomposition", index)
    add_page_header(page, "KSH Practice Capability Decomposition", "Current implementation plus approved Phase 13E-13H boundaries")
    root_id = page.vertex("<b>/practice</b><br><font color='#5F6B7A'>immutable attempts · skill-native delivery · evidence-based results</font>", 580, 135, 440, 90, status_box_style("CURRENT"))
    positions: dict[str, tuple[float, float]] = {}
    for i, cap in enumerate(CAPABILITIES):
        row = i // 5
        col = i % 5
        x = 55 + col * 305
        y = 300 + row * 280
        statuses = sorted({item["status"] for item in cap["use_cases"] if item["status"] != "CURRENT"})
        status = statuses[0] if statuses else "CURRENT"
        label = f"<b>{cap['code']}</b><br><b>{cap['short']}</b><br><font color='#5F6B7A'>{cap['purpose']}</font><br><br><b>{status}</b>"
        node = page.vertex(label, x, y, 270, 195, status_box_style(status))
        positions[cap["code"]] = (x + 135, y + 97)
        page.edge(root_id, node, "", f"endArrow=block;endFill=1;strokeColor=#{COLORS['line']};rounded=1;orthogonalLoop=1;jettySize=auto;")
    add_status_legend(page, 905)


def add_context_page(mxfile: ET.Element, index: int) -> None:
    page = DrawioPage(mxfile, "01 System Context and Status", index)
    add_page_header(page, "Practice System Context", "Actors, internal capability boundary and explicit future/deferred work")
    page.vertex("<b>Người học</b><br>catalog · player · result · progress", 70, 230, 220, 95, "shape=umlActor;verticalLabelPosition=bottom;verticalAlign=top;html=1;fontSize=13;fontColor=#0B2545;")
    page.vertex("<b>Giảng viên / Reviewer</b><br>authoring · import · publish · recovery", 70, 510, 220, 110, "shape=umlActor;verticalLabelPosition=bottom;verticalAlign=top;html=1;fontSize=13;fontColor=#0B2545;")
    boundary = page.vertex("", 365, 145, 820, 680, f"rounded=1;whiteSpace=wrap;html=1;fillColor=#{COLORS['panel']};strokeColor=#{COLORS['blue']};strokeWidth=2;arcSize=6;")
    page.vertex("<b>KSH PRACTICE</b>", 390, 160, 300, 34, drawio_text_style(17, COLORS["blue_dark"], True), parent=boundary)
    for i, cap in enumerate(CAPABILITIES):
        row = i // 2
        col = i % 2
        x = 395 + col * 385
        y = 215 + row * 112
        statuses = sorted({item["status"] for item in cap["use_cases"] if item["status"] != "CURRENT"})
        status = statuses[0] if statuses else "CURRENT"
        page.vertex(f"<b>{cap['code']}</b> · {cap['short']}<br><font color='#5F6B7A'>{status}</font>", x, y, 350, 78, status_box_style(status), parent=boundary)
    page.vertex("<b>AI providers</b><br>PDF · R/L explanation · Writing · Speaking", 1280, 210, 240, 100, status_box_style("EXTERNAL"))
    page.vertex("<b>Private media / DB</b><br>version bindings · attempts · audio", 1280, 445, 240, 100, status_box_style("EXTERNAL"))
    page.vertex("<b>Deferred gates</b><br>13H browser/device/a11y/perf<br>15 compatibility cleanup + premium seed", 1260, 675, 270, 125, status_box_style("DEFERRED 13H"))
    page.absolute_edge(290, 280, 365, 280, "learner actions", f"endArrow=block;strokeColor=#{COLORS['blue']};fontSize=11;")
    page.absolute_edge(290, 560, 365, 560, "authoring/review", f"endArrow=block;strokeColor=#{COLORS['blue']};fontSize=11;")
    page.absolute_edge(1185, 270, 1280, 270, "audited calls", f"endArrow=block;strokeColor=#{COLORS['warning']};fontSize=11;")
    page.absolute_edge(1185, 500, 1280, 500, "authorized data", f"endArrow=block;strokeColor=#{COLORS['blue_dark']};fontSize=11;")
    add_status_legend(page, 905)


def add_use_case_page(mxfile: ET.Element, cap: dict[str, Any], index: int) -> None:
    page = DrawioPage(mxfile, f"{cap['code']}-UC Use Cases", index)
    add_page_header(page, f"{cap['code']} · {cap['name']}", cap["purpose"])
    actor_ids = []
    for i, actor in enumerate(cap["actors"]):
        actor_ids.append(page.vertex(actor, 55, 230 + i * 250, 210, 95, "shape=umlActor;verticalLabelPosition=bottom;verticalAlign=top;html=1;fontSize=13;fontColor=#0B2545;"))
    boundary = page.vertex("", 335, 155, 850, 680, f"rounded=1;whiteSpace=wrap;html=1;fillColor=#{COLORS['panel']};strokeColor=#{COLORS['blue']};strokeWidth=2;arcSize=5;")
    page.vertex(f"<b>{cap['short']}</b>", 360, 170, 400, 36, drawio_text_style(16, COLORS["blue_dark"], True), parent=boundary)
    uc_ids = []
    for i, item in enumerate(cap["use_cases"]):
        fill, stroke = STATUS_STYLE[item["status"]]
        label = f"<b>{item['id']}</b><br>{item['title']}<br><font color='#{stroke}'><b>{item['status']}</b></font>"
        uc_id = page.vertex(label, 485, 245 + i * 180, 520, 115, f"ellipse;whiteSpace=wrap;html=1;fillColor=#{fill};strokeColor=#{stroke};fontColor=#{COLORS['ink']};fontSize=13;spacing=8;", parent=boundary)
        uc_ids.append(uc_id)
        page.edge(actor_ids[min(i, len(actor_ids) - 1)], uc_id, "", f"endArrow=none;strokeColor=#{COLORS['blue_dark']};")
    for i, secondary in enumerate(cap["secondary"]):
        ext = page.vertex(secondary, 1290, 255 + i * 255, 220, 100, status_box_style("EXTERNAL"))
        page.edge(uc_ids[min(i, len(uc_ids) - 1)], ext, "uses", f"endArrow=open;dashed=1;strokeColor=#{COLORS['external_stroke']};fontSize=10;")
    add_status_legend(page, 905)


def class_html(name: str, stereotype: str, status: str, responsibilities: list[str]) -> str:
    fill, stroke = STATUS_STYLE[status]
    details = "<br>".join(f"· {line}" for line in responsibilities)
    return f"<font color='#{stroke}'><b>{status}</b></font><br><i>«{stereotype}»</i><br><b>{name}</b><hr>{details}"


def add_class_page(mxfile: ET.Element, cap: dict[str, Any], index: int) -> None:
    page = DrawioPage(mxfile, f"{cap['code']}-CL Class Diagram", index)
    add_page_header(page, f"{cap['code']} · Class Diagram", "Current classes and explicitly labelled logical targets")
    nodes: dict[str, tuple[str, float, float]] = {}
    for i, (name, stereotype, status, responsibilities) in enumerate(cap["classes"]):
        row = i // 4
        col = i % 4
        x = 55 + col * 385
        y = 155 + row * 330
        cell_id = page.vertex(class_html(name, stereotype, status, responsibilities), x, y, 345, 235, status_box_style(status, rounded=False))
        nodes[name] = (cell_id, x + 172.5, y + 117.5)
    for source, target, label in cap["relations"]:
        if source in nodes and target in nodes:
            page.edge(nodes[source][0], nodes[target][0], label, f"endArrow=open;strokeColor=#{COLORS['blue_dark']};fontSize=10;rounded=1;orthogonalLoop=1;jettySize=auto;")
    add_status_legend(page, 905)


def add_sequence_page(mxfile: ET.Element, cap: dict[str, Any], item: dict[str, Any], sequence_number: int, index: int) -> None:
    page = DrawioPage(mxfile, f"{cap['code']}-S{sequence_number:02d} {item['id']}", index)
    add_page_header(page, f"{item['id']} · {item['title']}", f"{cap['short']} · {item['status']} · normal flow with fail-safe return")
    participants = item["participants"]
    count = len(participants)
    left = 60
    right = 1540
    slot = (right - left) / count
    centers: dict[str, float] = {}
    for i, participant in enumerate(participants):
        x_center = left + slot * i + slot / 2
        centers[participant] = x_center
        status = "CURRENT"
        if "[13E]" in participant:
            status = "PLANNED 13E"
        elif "[13F]" in participant:
            status = "PLANNED 13F"
        elif participant in ("AI Provider", "Transcription Provider", "Evaluation Provider", "Private Storage", "Browser Media API", "Database"):
            status = "EXTERNAL"
        width = max(130, min(205, slot - 18))
        page.vertex(participant, x_center - width / 2, 145, width, 58, status_box_style(status))
        page.absolute_edge(x_center, 205, x_center, 850, "", f"endArrow=none;dashed=1;strokeColor=#{COLORS['line']};")
    y = 240
    for number, (source, target, label, kind) in enumerate(item["sequence"], start=1):
        x1 = centers[source]
        x2 = centers[target]
        line_style = f"endArrow=block;endFill=1;strokeColor=#{COLORS['blue_dark']};fontSize=10;"
        if kind == "return":
            line_style = f"endArrow=open;dashed=1;strokeColor=#{COLORS['muted']};fontSize=10;"
        page.absolute_edge(x1, y, x2, y, f"{number}. {label}", line_style)
        y += 68
    alt_y = min(855, y + 10)
    page.vertex("alt  validation / authorization / provider failure", 70, alt_y, 1460, 42, f"rounded=0;whiteSpace=wrap;html=1;fillColor=#{COLORS['panel']};strokeColor=#{COLORS['warning']};dashed=1;fontColor=#{COLORS['warning']};fontSize=11;fontStyle=1;align=left;spacingLeft=10;")
    add_status_legend(page, 920)


def generate_drawio() -> None:
    expected_page_count = 2 + len(CAPABILITIES) * 2 + sum(len(cap["use_cases"]) for cap in CAPABILITIES)
    mxfile = ET.Element(
        "mxfile",
        {
            "host": "app.diagrams.net",
            "modified": DRAWIO_DATE,
            "agent": "Codex Practice Architecture Generator",
            "version": "26.0.9",
            "type": "device",
            "compressed": "false",
            "pages": str(expected_page_count),
        },
    )
    page_index = 0
    add_feature_decomposition(mxfile, page_index)
    page_index += 1
    add_context_page(mxfile, page_index)
    page_index += 1
    for cap in CAPABILITIES:
        add_use_case_page(mxfile, cap, page_index)
        page_index += 1
    for cap in CAPABILITIES:
        add_class_page(mxfile, cap, page_index)
        page_index += 1
    for cap in CAPABILITIES:
        for sequence_number, item in enumerate(cap["use_cases"], start=1):
            add_sequence_page(mxfile, cap, item, sequence_number, page_index)
            page_index += 1
    if page_index != expected_page_count:
        raise ValueError(f"Expected {expected_page_count} model-derived pages, generated {page_index}")
    ET.indent(mxfile, space="  ")
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    tree = ET.ElementTree(mxfile)
    tree.write(DRAWIO_PATH, encoding="utf-8", xml_declaration=True)


def wrap_preview(text: str, width: int) -> list[str]:
    plain = text.replace("<br>", "\n")
    while "<" in plain and ">" in plain:
        start = plain.find("<")
        end = plain.find(">", start)
        if end < 0:
            break
        plain = plain[:start] + plain[end + 1 :]
    lines: list[str] = []
    for part in plain.splitlines() or [plain]:
        lines.extend(textwrap.wrap(part, width=max(12, width)) or [""])
    return lines


def generate_preview_contact_sheets(preview_dir: Path) -> None:
    """Render lightweight PNG QA sheets from the same capability model.

    These are internal review aids, not deliverables and not a replacement for
    opening the valid Draw.io XML in diagrams.net.
    """
    from PIL import Image, ImageDraw, ImageFont

    preview_dir.mkdir(parents=True, exist_ok=True)
    try:
        font_regular = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 18)
        font_small = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 14)
        font_label = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 17)
        font_bold = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 20)
        font_title = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 28)
    except OSError:
        font_regular = font_small = font_label = font_bold = font_title = ImageFont.load_default()

    def color(hex_value: str) -> tuple[int, int, int]:
        return tuple(int(hex_value[i : i + 2], 16) for i in (0, 2, 4))

    def canvas(title: str, subtitle: str):
        image = Image.new("RGB", (1600, 1000), "white")
        draw = ImageDraw.Draw(image)
        draw.text((55, 30), title, fill=color(COLORS["ink"]), font=font_title)
        draw.text((55, 72), subtitle, fill=color(COLORS["muted"]), font=font_small)
        draw.line((55, 112, 1545, 112), fill=color(COLORS["line"]), width=2)
        return image, draw

    # Capability overview preview.
    image, draw = canvas("KSH Practice Capability Decomposition", "QA preview generated from the same model as Draw.io")
    draw.rounded_rectangle((580, 135, 1020, 225), radius=10, fill=color(COLORS["current_fill"]), outline=color(COLORS["current_stroke"]), width=3)
    draw.text((705, 165), "/practice", fill=color(COLORS["ink"]), font=font_bold)
    for i, cap in enumerate(CAPABILITIES):
        row, col = divmod(i, 5)
        x = 55 + col * 305
        y = 300 + row * 280
        statuses = sorted({item["status"] for item in cap["use_cases"] if item["status"] != "CURRENT"})
        status = statuses[0] if statuses else "CURRENT"
        fill, stroke = STATUS_STYLE[status]
        draw.rounded_rectangle((x, y, x + 270, y + 195), radius=8, fill=color(fill), outline=color(stroke), width=3)
        for line_no, line in enumerate(textwrap.wrap(f"{cap['code']} · {cap['short']}", width=27)[:2]):
            draw.text((x + 12, y + 12 + line_no * 19), line, fill=color(COLORS["ink"]), font=font_label)
        for line_no, line in enumerate(textwrap.wrap(cap["purpose"], width=34)[:5]):
            draw.text((x + 12, y + 58 + line_no * 20), line, fill=color(COLORS["muted"]), font=font_small)
        draw.text((x + 12, y + 169), status, fill=color(stroke), font=font_small)
    image.save(preview_dir / "00-feature-decomposition.png")

    for cap in CAPABILITIES:
        image, draw = canvas(f"{cap['code']} · {cap['short']}", "Use Case, class boundary and sequence traceability preview")
        draw.rounded_rectangle((55, 145, 520, 850), radius=8, fill=color(COLORS["panel"]), outline=color(COLORS["blue"]), width=3)
        draw.text((75, 165), "USE CASES", fill=color(COLORS["blue_dark"]), font=font_bold)
        for i, item in enumerate(cap["use_cases"]):
            y = 225 + i * 190
            fill, stroke = STATUS_STYLE[item["status"]]
            draw.rounded_rectangle((85, y, 490, y + 135), radius=60, fill=color(fill), outline=color(stroke), width=3)
            draw.text((115, y + 18), item["id"], fill=color(stroke), font=font_bold)
            for line_no, line in enumerate(textwrap.wrap(item["title"], width=38)[:3]):
                draw.text((115, y + 55 + line_no * 21), line, fill=color(COLORS["ink"]), font=font_regular)
        draw.rounded_rectangle((560, 145, 1540, 850), radius=8, fill="white", outline=color(COLORS["line"]), width=2)
        draw.text((585, 165), "CLASS BOUNDARY", fill=color(COLORS["blue_dark"]), font=font_bold)
        for i, (name, stereotype, status, responsibilities) in enumerate(cap["classes"]):
            row, col = divmod(i, 4)
            x = 585 + col * 235
            y = 225 + row * 295
            fill, stroke = STATUS_STYLE[status]
            draw.rounded_rectangle((x, y, x + 210, y + 220), radius=5, fill=color(fill), outline=color(stroke), width=2)
            draw.text((x + 9, y + 9), status, fill=color(stroke), font=font_small)
            draw.text((x + 9, y + 36), stereotype, fill=color(COLORS["muted"]), font=font_small)
            for line_no, line in enumerate(textwrap.wrap(name, width=22)[:3]):
                draw.text((x + 9, y + 62 + line_no * 21), line, fill=color(COLORS["ink"]), font=font_bold)
            for line_no, line in enumerate(responsibilities[:2]):
                for wrap_no, wrapped in enumerate(textwrap.wrap(line, width=24)[:2]):
                    draw.text((x + 9, y + 132 + (line_no * 2 + wrap_no) * 18), wrapped, fill=color(COLORS["muted"]), font=font_small)
        image.save(preview_dir / f"{cap['code'].lower()}-capability.png")

    # Sequence contact sheets: one mini-tile per Use Case.
    for cap in CAPABILITIES:
        image, draw = canvas(f"{cap['code']} · Sequence Diagrams", "Three formal flows; dashed arrows are returns")
        for tile_index, item in enumerate(cap["use_cases"]):
            top = 145 + tile_index * 265
            draw.rounded_rectangle((55, top, 1545, top + 230), radius=6, fill="white", outline=color(COLORS["line"]), width=2)
            draw.text((70, top + 10), f"{item['id']} · {item['title']}", fill=color(COLORS["ink"]), font=font_bold)
            participants = item["participants"]
            slot = 1420 / len(participants)
            centers = {}
            for i, participant in enumerate(participants):
                center = 85 + i * slot + slot / 2
                centers[participant] = center
                draw.text((center - slot * 0.42, top + 48), participant[:24], fill=color(COLORS["blue_dark"]), font=font_small)
                draw.line((center, top + 72, center, top + 215), fill=color(COLORS["line"]), width=1)
            y = top + 88
            for source, target, label, kind in item["sequence"][:6]:
                x1, x2 = centers[source], centers[target]
                line_color = COLORS["muted"] if kind == "return" else COLORS["blue_dark"]
                draw.line((x1, y, x2, y), fill=color(line_color), width=2)
                arrow_x = x2
                direction = 1 if x2 > x1 else -1
                draw.polygon([(arrow_x, y), (arrow_x - direction * 8, y - 4), (arrow_x - direction * 8, y + 4)], fill=color(line_color))
                draw.text((min(x1, x2) + 4, y - 16), label[:44], fill=color(line_color), font=font_small)
                y += 22
        image.save(preview_dir / f"{cap['code'].lower()}-sequences.png")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--preview-dir", type=Path, help="Optional internal PNG QA directory")
    output_group = parser.add_mutually_exclusive_group()
    output_group.add_argument("--docx-only", action="store_true", help="Generate only the DOCX artifact")
    output_group.add_argument("--drawio-only", action="store_true", help="Generate only the Draw.io artifact")
    args = parser.parse_args()
    validate_model()
    if not args.drawio_only:
        generate_docx()
    if not args.docx_only:
        generate_drawio()
    if args.preview_dir:
        generate_preview_contact_sheets(args.preview_dir)
    if not args.drawio_only:
        print(f"Generated {DOCX_PATH}")
    if not args.docx_only:
        print(f"Generated {DRAWIO_PATH}")
    if args.preview_dir:
        print(f"Generated previews in {args.preview_dir}")


if __name__ == "__main__":
    main()
